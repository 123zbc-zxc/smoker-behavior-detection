from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
BACKUP_PATH = ROOT / "output" / "doc" / "thesis_backup_before_ai_risk_rewrite_round3.docx"
ALIGNED_COPY_PATH = ROOT / "output" / "doc" / "thesis_sample_aligned.docx"


REPLACEMENTS: dict[int, str] = {
    42: (
        "摘要：在医院、校园、办公楼和交通站点等禁烟场景中，"
        "吸烟行为管理长期依赖人工巡查和人工回看监控。"
        "这种方式在场景较少时尚能维持，但监控点位一多，就容易出现覆盖不全、发现滞后和维护成本偏高的问题。"
        "基于这一现实需求，本文围绕吸烟行为检测开展数据整理、模型训练和展示系统实现。"
        "本文以 YOLOv8n 为基础模型，将检测目标统一为 cigarette、smoking_person 和 smoke 三类。"
        "针对 AI Studio、Roboflow 以及新增压缩包数据在格式、类别和标注规则上的差异，"
        "完成了数据转换、标签映射、异常标签清理、数据合并和数据集划分，构建出可直接训练的吸烟行为检测数据集。"
        "考虑到本地实验环境以 CPU 为主，若直接在全量数据集上长时间反复训练，实验推进会明显受限，"
        "因此本文又构建了温和均衡版数据集 smoke_bal，作为当前阶段的主实验集。"
        "在模型部分，本文先完成 YOLOv8n 的训练、验证、推理与导出流程，再在颈部特征融合位置引入 ECA 通道注意力模块，"
        "形成 YOLOv8n+ECA 对比模型。实验结果表明，baseline 模型在 smoke_bal 测试集上的 Precision、Recall、mAP@0.5 和 "
        "mAP@0.5:0.95 分别为 0.526、0.625、0.520 和 0.323；ECA 模型对应结果为 0.513、0.572、0.494 和 0.299。"
        "从结果看，当前 baseline 表现更稳定，smoke 类识别效果较好，而 cigarette 仍然是本课题中最突出的难点。"
        "在系统实现方面，项目已经完成训练、验证、预测和模型导出脚本，并搭建了基于 FastAPI 的网页展示方案，"
        "用于集中呈现实验流程、检测结果和案例演示。"
    ),
    48: (
        "Abstract: This work focuses on smoker behavior detection in smoke-free public areas. "
        "In practical management, manual inspection and manual video review are still common, but they become inefficient "
        "once the number of cameras grows. For this reason, the thesis studies an automatic detection pipeline based on deep learning."
    ),
    49: (
        "YOLOv8n is used as the baseline detector, and the task is unified into three categories: cigarette, smoking_person, and smoke. "
        "Data collected from AI Studio, Roboflow, and added ZIP datasets are converted, remapped, cleaned, merged, and split into a trainable YOLO dataset. "
        "Because most local experiments are conducted on CPU, a mildly balanced dataset named smoke_bal is further built for the main comparison experiments."
    ),
    50: (
        "On smoke_bal, the baseline model reaches Precision 0.526, Recall 0.625, mAP@0.5 0.520, and mAP@0.5:0.95 0.323 after 30 epochs. "
        "An ECA-based variant is also trained under the same setting, but it does not outperform the baseline in the current budget. "
        "The project also completes scripts for training, validation, prediction, and model export, and builds a FastAPI-based web demo for defense presentation and result display."
    ),
    57: (
        "吸烟不仅危害吸烟者本人的健康，也会对周围人群造成二手烟影响。"
        "在医院、教学楼、宿舍楼、办公区域和交通枢纽等公共场所，禁烟管理一直都是实际工作中的常见问题。"
        "虽然这些区域大多已经安装监控设备，但日常管理仍然主要依赖人工巡查或事后回看监控。"
        "这种方式在点位较少时还能维持，一旦场景扩大、时间跨度变长，就容易出现顾不过来、发现不及时等情况。"
    ),
    58: (
        "对禁烟场景来说，真正需要的并不只是事后追溯，更重要的是尽早发现并及时处理。"
        "如果能够直接利用现有监控画面识别吸烟行为，就可以在一定程度上减轻人工压力，"
        "也便于后续扩展到告警提示、截图留存和记录统计等功能。"
        "因此，吸烟行为检测并不是脱离实际的研究题目，而是与公共场所管理需求联系紧密的应用任务。"
    ),
    59: (
        "不过，这个任务本身并不好做。"
        "首先，香烟在图像里通常属于小目标，尺寸很小，特征也不明显；"
        "其次，吸烟动作会受到拍摄角度、人物姿态和遮挡关系影响；"
        "再次，烟雾形态不固定，边界又往往不清楚，容易受到背景和光照干扰。"
        "也就是说，吸烟行为检测同时包含了小目标检测和复杂场景识别两方面难点。"
        "正因如此，本文围绕这一任务开展数据整理、模型训练和系统展示设计，希望形成一条完整且可落地的毕业设计实现链路。"
    ),
    61: (
        "目标检测的发展经历了比较清楚的技术演变。"
        "早期方法更多依赖人工设计特征，再配合分类器完成识别，例如 Haar、HOG、SIFT 等。"
        "这类方法在背景简单、规则比较固定的图像中还能发挥作用，"
        "但遇到目标尺度变化、局部遮挡或复杂背景时，稳定性往往下降得很快。"
    ),
    62: (
        "深度学习兴起后，目标检测的主流路线逐渐转向端到端学习。"
        "两阶段方法如 Faster R-CNN 通常在精度上更有优势，但推理速度相对偏慢；"
        "单阶段方法如 SSD 和 YOLO 系列则更强调检测速度与部署效率之间的平衡。"
        "对本课题这种既要完成训练实验、又要接入展示系统的任务来说，后者更符合实际条件。"
    ),
    63: (
        "具体到吸烟行为检测，研究难点通常集中在两个方面。"
        "一是怎样把 cigarette 这类小目标尽量稳定地检出来，"
        "二是怎样让模型在监控画面这种复杂背景下保持相对稳定。"
        "围绕这些问题，已有研究常见的处理思路包括改进特征融合、引入注意力机制、调整样本分布，"
        "或者针对小目标单独优化输入尺度与训练策略。"
    ),
    64: (
        "结合本项目现有的数据基础和代码实现，本文并没有把重点放在追求某一组参数的极限结果上，"
        "而是更关心数据整理、模型训练、结果验证和系统展示能否真正接起来。"
        "也正因为如此，后文的写法会更偏向工程实现与结果分析，而不是只讨论网络结构本身。"
    ),
    69: "围绕上述问题，本文主要完成了以下几项工作：",
    70: "1. 整理 AI Studio、Roboflow 以及新增压缩包数据，并统一为 cigarette、smoking_person 和 smoke 三类检测任务；",
    71: "2. 编写并完善数据处理脚本，完成格式转换、标签映射、异常标签清理、数据集划分和完整性检查；",
    72: "3. 搭建 YOLOv8n baseline 的训练、验证、推理和导出流程，形成基础实验框架；",
    73: "4. 在 YOLOv8n 颈部特征融合位置引入 ECA 模块，构建 YOLOv8n+ECA 对比模型；",
    74: "5. 结合答辩展示需求，完成基于 FastAPI 的网页展示方案，用于统一呈现实验流程、检测结果与系统状态。",
    82: (
        "全文共分为六章。"
        "第一章说明课题背景、研究现状和本文工作；"
        "第二章介绍卷积神经网络、YOLOv8n、ECA 注意力机制以及展示部署相关技术；"
        "第三章说明数据集构建过程与训练方案设计；"
        "第四章给出实验环境、评价指标及实验结果分析；"
        "第五章介绍系统实现与展示方案；"
        "第六章对全文工作进行总结，并说明后续仍需继续完善的部分。"
    ),
    86: (
        "在本课题中，卷积神经网络主要承担特征提取任务。"
        "靠近输入端的卷积层更容易保留边缘、亮度变化和纹理信息，"
        "随着网络层数加深，特征会逐步从局部细节过渡到更抽象的语义表达。"
        "和早期依赖人工挑选特征的做法相比，这种方式的优势在于模型能够在训练过程中自己学习哪些信息更有助于识别。"
    ),
    87: (
        "目标检测与普通图像分类的差别在于，它不仅要判断画面里有什么，还要给出目标所在位置。"
        "对吸烟行为检测来说，这一差别尤其关键。"
        "人物主体通常相对容易识别，烟雾次之，真正困难的是面积较小的香烟本体。"
        "因此，网络能否在多尺度特征之间保留足够的细粒度信息，会直接影响后续实验结果。"
    ),
    89: (
        "YOLO 系列属于单阶段检测框架，推理时一次前向传播就能同时完成类别预测和位置回归。"
        "对本科毕业设计这样的课题来说，它的优势比较实际：检测链路短，运行速度快，也更方便和后续网页系统对接。"
    ),
    90: (
        "结合当前硬件条件，本文没有继续扩大模型规模，而是把 YOLOv8n 作为 baseline。"
        "一方面，本地训练长期依赖 CPU，如果直接换成更大的 YOLOv8s 或 YOLOv8m，实验轮次和对比效率都会明显下降；"
        "另一方面，后续网页展示也需要调用同一套权重进行推理，因此轻量模型更符合本课题现阶段的边界。"
        "从结构上看，YOLOv8n 通过 C2f 模块提取主干特征，在颈部完成多尺度融合，再由 Anchor-Free 检测头输出类别与边界框。"
    ),
    96: (
        "注意力机制的作用并不是简单加一个“更强模块”，"
        "而是在特征传递过程中重新分配关注重点。"
        "ECA 属于通道注意力方法，它没有使用复杂的降维和升维结构，"
        "而是直接通过局部一维卷积处理相邻通道之间的关系，因此额外参数量较小，比较适合插入轻量网络。"
    ),
    97: (
        "放到本课题中看，ECA 的意义主要在于尝试强化弱特征。"
        "香烟目标面积小、纹理细，很容易被背景信息淹没；"
        "如果特征融合阶段对关键通道的响应不够敏感，后面的检测头就难以拿到稳定信息。"
        "基于这一考虑，本文才在后续模型设计中尝试把 ECA 接入 YOLOv8n。"
    ),
    101: (
        "数据增强在本课题中并不是为了机械增加样本数量，"
        "而是为了让模型提前见到更接近真实监控画面的变化形式。"
        "亮度波动、轻微模糊、尺度变化和拍摄角度偏移，都可能出现在实际场景里。"
        "适度增强的作用，是降低模型对单一采集条件的依赖。"
    ),
    102: (
        "小目标检测的难点则更直接。"
        "目标一旦过小，下采样后能保留下来的边界和纹理信息就会快速减少，"
        "检测头可利用的判别依据也会随之变弱。"
        "围绕这一问题，常见办法包括提高输入分辨率、加强浅层细节特征参与、改善样本分布以及引入轻量注意力模块。"
        "本文第三章中的数据均衡、baseline 参数设置和 ECA 结构尝试，都是围绕这一难点展开的。"
    ),
    104: (
        "本课题最终采用网页展示，而没有停留在脚本级演示，"
        "主要是出于答辩展示和后续扩展两方面考虑。"
        "浏览器入口更统一，也更便于把实验摘要、模型状态和检测结果集中放在同一界面中。"
        "FastAPI 与现有 Python 推理代码的衔接较自然，因此被选作当前原型系统的后端框架。"
    ),
    105: (
        "在具体实现中，页面模板负责组织展示内容，静态资源目录负责管理上传文件和结果文件，"
        "数据库负责保存模型配置、检测记录和任务状态。"
        "考虑到当前系统主要用于单机演示，SQLite 已经能够满足现阶段需求；"
        "如果后续扩展到长期运行或多人使用，再切换到 PostgreSQL 会更合适。"
    ),
    106: (
        "模型导出部分主要为后续部署预留接口。"
        "当前训练和网页推理仍以 PyTorch 权重为主，但项目同时保留了 ONNX 导出脚本。"
        "这样一来，如果后续需要接入其他推理环境，就不必再从头整理模型转换流程。"
    ),
    108: (
        "本章没有单独铺陈过多理论，而是把后续实验真正会用到的技术背景先交代清楚，"
        "包括检测网络、ECA 机制、小目标问题以及网页系统选型。"
        "这些内容分别对应第三章的数据处理、第四章的实验分析和第五章的系统实现。"
    ),
}


def find_target_docx() -> Path:
    matches = [
        p
        for p in ROOT.glob("*202212903403*.docx")
        if "定稿" in p.name and not p.name.startswith("~$")
    ]
    if not matches:
        raise FileNotFoundError("未找到当前定稿论文 DOCX 文件。")
    matches.sort()
    return matches[0]


def backup_docx(target: Path) -> None:
    BACKUP_PATH.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(target, BACKUP_PATH)


def rewrite_doc(doc_path: Path) -> None:
    doc = Document(doc_path)
    for idx, text in REPLACEMENTS.items():
        if idx >= len(doc.paragraphs):
            raise IndexError(f"Paragraph index out of range: {idx}")
        doc.paragraphs[idx].text = text
    doc.save(doc_path)


def main() -> None:
    target = find_target_docx()
    backup_docx(target)
    rewrite_doc(target)
    if ALIGNED_COPY_PATH.exists():
        shutil.copy2(target, ALIGNED_COPY_PATH)
    print(f"Backup saved to: {BACKUP_PATH}")
    print(f"Round-3 rewritten thesis saved to: {target}")
    if ALIGNED_COPY_PATH.exists():
        print(f"Aligned copy synced to: {ALIGNED_COPY_PATH}")


if __name__ == "__main__":
    main()
