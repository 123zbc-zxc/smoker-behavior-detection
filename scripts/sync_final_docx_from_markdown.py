from __future__ import annotations

from pathlib import Path
import sys

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

ROOT = Path(__file__).resolve().parents[1]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

from scripts.build_standardized_thesis_docx import parse_blocks


SOURCE_MD = ROOT / "docs" / "paper" / "graduation_final_draft.md"
BACKUP_PATH = ROOT / "output" / "doc" / "final_docx_backup_before_sync.docx"
FALLBACK_OUTPUT_PATH = ROOT / "output" / "doc" / "thesis_final_synced_copy.docx"

START_CHAPTER_TITLE = "第2章 相关理论与关键技术"
REFERENCES_TITLE = "参考文献"
APPENDIX_TITLE = "附录：项目对应的关键命令"

IMAGE_PLAN: dict[str, list[tuple[str, str, float]]] = {
    "2.2 YOLO系列模型与YOLOv8n结构特点": [
        ("figures/图6_YOLOv8各版本性能对比表.png", "图2-1 YOLOv8各版本性能对比表", 14.5),
        ("figures/图7_YOLOv8n网络结构图.png", "图2-2 YOLOv8n网络结构图", 14.5),
    ],
    "2.3 ECA注意力机制原理": [
        ("figures/图8_ECA-Net模块结构示意图.png", "图2-3 ECA-Net模块结构示意图", 12.8),
    ],
    "4.4 分类别结果对比分析": [
        ("runs/val/smoking_eval/confusion_matrix_normalized.png", "图4-1 Baseline模型测试集归一化混淆矩阵", 12.8),
    ],
    "4.5 训练过程与收敛情况分析": [
        ("runs/train/yolov8n_balanced_30/results.png", "图4-2 Baseline模型训练结果曲线", 14.0),
    ],
    "5.1 系统总体架构": [
        ("figures/图13_系统总体架构图.png", "图5-2 系统总体架构图", 15.0),
    ],
    "5.4 FastAPI 前端展示设计": [
        ("output/web_demo/results/images/bcb33fe4b2f94c0387e08ffdb01184a7.jpg", "图5-1 网页端图片检测结果示意图", 14.0),
    ],
    "5.6 数据库与历史记录模块设计": [
        ("figures/图14_数据库关系说明图.png", "图5-3 数据库关系说明图", 15.0),
    ],
}


def find_target_docx() -> Path:
    candidates = [p for p in ROOT.glob("定稿_智能224-*.docx") if not p.name.startswith("~$")]
    if not candidates:
        raise FileNotFoundError("未找到最终定稿 DOCX 文件。")
    return candidates[0]


def backup_docx(target: Path) -> Path:
    BACKUP_PATH.parent.mkdir(parents=True, exist_ok=True)
    BACKUP_PATH.write_bytes(target.read_bytes())
    return BACKUP_PATH


def set_update_fields(doc: Document) -> None:
    settings = doc.settings.element
    found = settings.find(qn("w:updateFields"))
    if found is None:
        update_fields = OxmlElement("w:updateFields")
        update_fields.set(qn("w:val"), "true")
        settings.append(update_fields)
    else:
        found.set(qn("w:val"), "true")


def set_run_font(
    run,
    *,
    east_asia: str = "宋体",
    latin: str = "Times New Roman",
    size: float = 12,
    bold: bool = False,
) -> None:
    run.font.name = latin
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), east_asia)
    run.font.size = Pt(size)
    run.bold = bold


def body_text(element) -> str:
    return "".join(node.text or "" for node in element.iter() if node.tag == qn("w:t")).strip()


def locate_replace_child_index(doc: Document) -> int:
    for idx, child in enumerate(doc._element.body.iterchildren()):
        if child.tag == qn("w:p") and body_text(child).startswith(START_CHAPTER_TITLE):
            return idx
    raise ValueError("未能在最终定稿中定位到第2章起始位置。")


def truncate_document_from(doc: Document, start_child_index: int) -> None:
    body = doc._element.body
    children = list(body.iterchildren())
    for child in children[start_child_index:]:
        if child.tag == qn("w:sectPr"):
            continue
        body.remove(child)


def collect_markdown_blocks() -> list[tuple[str, object]]:
    blocks = parse_blocks(SOURCE_MD.read_text(encoding="utf-8"))
    selected: list[tuple[str, object]] = []
    started = False
    for kind, payload in blocks:
        if kind == "heading":
            level, text, _raw = payload
            if text == START_CHAPTER_TITLE:
                started = True
            if not started:
                continue
            if text == "Abstract":
                continue
            if level == 1:
                continue
        if started:
            selected.append((kind, payload))
    return selected


def add_heading(doc: Document, text: str, style: str) -> None:
    paragraph = doc.add_paragraph(style=style)
    paragraph.paragraph_format.first_line_indent = Cm(0)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    if style == "Heading 1":
        set_run_font(run, east_asia="黑体", size=16, bold=True)
    else:
        set_run_font(run, east_asia="黑体", size=14, bold=True)


def add_normal_paragraph(doc: Document, text: str, *, first_line_indent: bool = True) -> None:
    paragraph = doc.add_paragraph(style="Normal")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.first_line_indent = Cm(0.74) if first_line_indent else Cm(0)
    run = paragraph.add_run(text)
    set_run_font(run, east_asia="宋体", size=12)


def add_code_paragraph(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="Normal")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.first_line_indent = Cm(0)
    run = paragraph.add_run(text)
    set_run_font(run, east_asia="宋体", latin="Courier New", size=10.5)


def add_table(doc: Document, table_lines: list[str]) -> None:
    rows: list[list[str]] = []
    for line in table_lines:
        stripped = line.strip("|")
        normalized = stripped.replace("|", "").replace("-", "").replace(":", "").strip()
        if not normalized:
            continue
        rows.append([cell.strip().replace("`", "") for cell in line.strip("|").split("|")])
    if not rows:
        return

    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            paragraph = table.cell(i, j).paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run(cell_text)
            set_run_font(run, east_asia="宋体", size=10.5, bold=(i == 0))
    doc.add_paragraph()


def add_image_with_caption(doc: Document, image_path: Path, caption: str, width_cm: float) -> None:
    if not image_path.exists():
        return
    paragraph = doc.add_paragraph(style="Normal")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Cm(0)
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Cm(width_cm))

    caption_paragraph = doc.add_paragraph(style="Normal")
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_paragraph.paragraph_format.first_line_indent = Cm(0)
    run = caption_paragraph.add_run(caption)
    set_run_font(run, east_asia="宋体", size=10.5)


def clean_inline_marks(text: str) -> str:
    return text.replace("`", "").replace("**", "")


def append_blocks(doc: Document, blocks: list[tuple[str, object]]) -> None:
    last_section_heading: str | None = None

    def flush_images(heading_text: str | None) -> None:
        if not heading_text:
            return
        for relative_path, caption, width_cm in IMAGE_PLAN.get(heading_text, []):
            add_image_with_caption(doc, ROOT / relative_path, caption, width_cm)

    for kind, payload in blocks:
        if kind == "heading":
            flush_images(last_section_heading)
            level, text, _raw = payload
            if text.startswith("第") and "章" in text:
                add_heading(doc, text, "Heading 1")
                last_section_heading = None
            elif text in {REFERENCES_TITLE, APPENDIX_TITLE}:
                add_heading(doc, text, "Heading 2")
                last_section_heading = None
            elif level >= 3 or text[:2].isdigit():
                add_heading(doc, text, "Heading 2")
                last_section_heading = text
            else:
                add_heading(doc, text, "Heading 2")
                last_section_heading = text
            continue

        if kind == "paragraph":
            text = clean_inline_marks(str(payload)).strip()
            if not text:
                doc.add_paragraph()
                continue
            add_normal_paragraph(doc, text)
            continue

        if kind == "table":
            add_table(doc, payload)
            continue

        if kind == "number":
            for index, item in enumerate(payload, start=1):
                add_normal_paragraph(doc, f"{index}. {clean_inline_marks(str(item))}", first_line_indent=False)
            continue

        if kind == "bullet":
            for item in payload:
                add_normal_paragraph(doc, f"• {clean_inline_marks(str(item))}", first_line_indent=False)
            continue

        if kind == "code":
            add_code_paragraph(doc, str(payload))

    flush_images(last_section_heading)


def sync_final_docx() -> tuple[Path, Path, Path]:
    target = find_target_docx()
    backup = backup_docx(target)
    doc = Document(target)
    set_update_fields(doc)
    start_index = locate_replace_child_index(doc)
    truncate_document_from(doc, start_index)
    append_blocks(doc, collect_markdown_blocks())
    try:
        doc.save(target)
        saved_path = target
    except PermissionError:
        FALLBACK_OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
        doc.save(FALLBACK_OUTPUT_PATH)
        saved_path = FALLBACK_OUTPUT_PATH
    return saved_path, backup, target


def main() -> None:
    saved_path, backup, target = sync_final_docx()
    if saved_path == target:
        print(f"Updated final thesis docx: {saved_path}")
    else:
        print(f"Target docx is locked, saved updated copy to: {saved_path}")
        print(f"Original target remains: {target}")
    print(f"Backup saved to: {backup}")


if __name__ == "__main__":
    main()
