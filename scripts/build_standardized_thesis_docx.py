from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[1]
SOURCE_MD = ROOT / 'docs' / 'paper' / 'graduation_final_draft.md'
OUTPUT_PATH = ROOT / 'output' / 'doc' / 'thesis_standardized_v3.docx'
FIG_DIR = ROOT / 'figures'
TEMPLATE_PATH = next(ROOT.glob('定稿*.docx'))

TITLE_CN = '基于深度学习的吸烟者行为检测系统设计与实现'
STUDENT_NAME = '陈刚'
COLLEGE = '新工科产业学院'
MAJOR = '智能科学与技术'
CLASS_NAME = '智能224'
STUDENT_ID = '202212903403'
ADVISOR = '林一锋'
ADVISOR_TITLE = '助教'
DATE_TEXT = '2026年4月'

IMAGE_PLAN = {
    '### 2.2 YOLOv8n 模型原理': [('图6_YOLOv8各版本性能对比表.png', '图2-1 YOLOv8各版本性能对比表'), ('图7_YOLOv8n网络结构图.png', '图2-2 YOLOv8n网络结构图')],
    '### 2.3 ECA 注意力机制': [('图8_ECA-Net模块结构示意图.png', '图2-3 ECA-Net模块结构示意图')],
    '### 3.5 YOLOv8n+ECA 改进模型设计': [('图9_YOLOv8n颈部特征融合结构示意图.png', '图3-1 YOLOv8n颈部特征融合结构示意图'), ('图12_YOLOv8n+ECA网络结构图.png', '图3-2 YOLOv8n+ECA网络结构图')],
}


def set_run_font(run, east_asia: str = '宋体', latin: str = 'Times New Roman', size: float = 12, bold: bool = False) -> None:
    run.font.name = latin
    run._element.rPr.rFonts.set(qn('w:eastAsia'), east_asia)
    run.font.size = Pt(size)
    run.bold = bold


def configure_styles(doc: Document, template: Document) -> None:
    tpl_sec = template.sections[0]
    sec = doc.sections[0]
    sec.page_width = tpl_sec.page_width
    sec.page_height = tpl_sec.page_height
    sec.top_margin = tpl_sec.top_margin
    sec.bottom_margin = tpl_sec.bottom_margin
    sec.left_margin = tpl_sec.left_margin
    sec.right_margin = tpl_sec.right_margin

    normal = doc.styles['Normal']
    normal.font.name = 'Times New Roman'
    normal._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    normal.font.size = Pt(12)
    pf = normal.paragraph_format
    pf.first_line_indent = Cm(0.74)
    pf.line_spacing = 1.5
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)

    style_map = {
        'Title': ('黑体', 18),
        'Heading 1': ('黑体', 16),
        'Heading 2': ('黑体', 14),
        'Heading 3': ('黑体', 12),
    }
    for name, (east, size) in style_map.items():
        style = doc.styles[name]
        style.font.name = 'Times New Roman'
        style._element.rPr.rFonts.set(qn('w:eastAsia'), east)
        style.font.size = Pt(size)
        style.font.bold = True


def clear_document(doc: Document) -> None:
    body = doc._element.body
    for child in list(body):
        body.remove(child)
    body.append(OxmlElement('w:sectPr'))


def add_center_line(doc: Document, text: str, size: float, bold: bool = False, east_asia: str = '黑体') -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run_font(run, east_asia=east_asia, size=size, bold=bold)


def add_page_break(doc: Document) -> None:
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def add_cover(doc: Document) -> None:
    add_center_line(doc, '本科毕业设计（论文）', 22, True)
    for _ in range(3):
        doc.add_paragraph()
    add_center_line(doc, TITLE_CN, 18, True)
    doc.add_paragraph()
    add_center_line(doc, STUDENT_NAME, 16, False)
    doc.add_paragraph()

    fields = [
        ('学    院：', COLLEGE),
        ('专    业：', MAJOR),
        ('班    级：', CLASS_NAME),
        ('学    号：', STUDENT_ID),
        ('指导教师：', ADVISOR),
        ('职    称：', ADVISOR_TITLE),
    ]
    for label, value in fields:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p.add_run(label + '    ')
        set_run_font(r1, east_asia='宋体', size=14)
        r2 = p.add_run(value)
        set_run_font(r2, east_asia='宋体', size=14)

    doc.add_paragraph()
    add_center_line(doc, DATE_TEXT, 14, False, east_asia='宋体')
    add_page_break(doc)

    add_center_line(doc, '原创性声明', 16, True)
    p = doc.add_paragraph()
    r = p.add_run('本人郑重声明：所提交的毕业设计（论文）是本人在指导教师指导下独立完成的研究成果。除文中已经注明引用的内容外，本文不包含任何其他个人或集体已经发表或撰写过的研究成果。对本文的研究做出重要贡献的个人和集体，均已在文中以明确方式标明。本人完全意识到本声明的法律后果由本人承担。')
    set_run_font(r)
    p = doc.add_paragraph()
    r = p.add_run('学生签名：____________________            日期：____________________')
    set_run_font(r, east_asia='宋体')

    add_center_line(doc, '指导声明', 16, True)
    p = doc.add_paragraph()
    r = p.add_run(f'本人指导的 {STUDENT_NAME} 同学毕业设计（论文）选题符合专业培养目标要求，工作量和难度适当。指导过程中已对论文内容进行了审阅与检查，确认其研究工作符合毕业设计基本规范，特此声明。')
    set_run_font(r)
    p = doc.add_paragraph()
    r = p.add_run('指导教师签名：____________________        日期：____________________')
    set_run_font(r, east_asia='宋体')
    add_page_break(doc)


def read_and_refine_markdown() -> str:
    md = SOURCE_MD.read_text(encoding='utf-8')
    md = md.replace('项目完成了数据集整理、标签映射、数据清洗、训练验证、推理导出和 PyQt5 可视化界面的完整闭环。', '项目完成了数据集整理、标签映射、数据清洗、训练验证、推理导出和展示层方案设计。其中，PyQt5 原型用于本地算法联调与快速验证，最终答辩展示将采用前端界面进行可视化呈现。')
    md = md.replace('**关键词**：吸烟行为检测；目标检测；YOLOv8n；ECA 注意力；深度学习；PyQt5', '**关键词**：吸烟行为检测；目标检测；YOLOv8n；ECA 注意力；深度学习；FastAPI')
    md = md.replace('5. 实现基于 PyQt5 的桌面检测演示系统，支持模型加载、图像检测和视频检测。', '5. 设计基于 FastAPI 的前端展示方案，并保留桌面原型用于本地联调与快速验证。')
    md = md.replace('### 2.4 PyQt5 与 ONNX 部署技术', '### 2.4 FastAPI 与 ONNX 展示部署技术')
    md = md.replace('PyQt5 能够快速构建桌面可视化界面，适合用于毕业设计答辩中的交互式展示。ONNX 则是常用的跨框架推理中间格式，可提升后续部署灵活性。本文在工程实现中同时保留 `.pt` 直接推理和 ONNX 导出能力：前者便于快速调试和演示，后者便于后续在 CPU 或其他推理框架中开展部署拓展。', 'FastAPI 适合在 Python 环境下快速搭建本地推理接口和展示服务，能够较自然地与当前 YOLOv8 推理脚本衔接，也更适合答辩场景下的网页化演示。ONNX 则是常用的跨框架推理中间格式，可提升后续部署灵活性。本文在工程实现中同时保留 `.pt` 直接推理、FastAPI 接口封装和 ONNX 导出能力：前者便于快速调试，FastAPI 便于构建前端展示页，后者便于后续在 CPU 或其他推理框架中开展部署拓展。')
    md = md.replace('| GUI 框架 | PyQt5 |', '| 展示后端 | FastAPI |')
    md = md.replace('从环境验证结果看，`ultralytics`、`torch` 和 `PyQt5` 已在当前虚拟环境中正常安装，`python scripts/train.py --help` 可正常运行，表明训练脚本与依赖环境已经完成联通。', '从环境验证结果看，`ultralytics` 与 `torch` 已在当前虚拟环境中正常安装，训练脚本可正常运行，说明模型训练与推理依赖已经完成联通。后续前端展示层将基于 FastAPI 进一步搭建。')
    md = md.replace('3. 将当前 PyQt5 原型与前端展示层结合，构建更适合答辩演示的可视化界面；', '3. 基于 FastAPI 完成前端展示层开发，构建更适合答辩演示的可视化界面；')
    return md


def parse_blocks(md: str) -> list[tuple[str, object]]:
    blocks: list[tuple[str, object]] = []
    lines = md.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        if not line.strip():
            i += 1
            continue
        if line.startswith('```'):
            code = []
            i += 1
            while i < len(lines) and not lines[i].startswith('```'):
                code.append(lines[i])
                i += 1
            blocks.append(('code', '\n'.join(code)))
            i += 1
            continue
        if re.match(r'^#{1,3}\s+', line):
            level = line.count('#', 0, line.find(' '))
            text = re.sub(r'^#{1,3}\s+', '', line).strip()
            blocks.append(('heading', (level, text, line)))
            i += 1
            continue
        if line.startswith('|'):
            table_lines = []
            while i < len(lines) and lines[i].startswith('|'):
                table_lines.append(lines[i].strip())
                i += 1
            blocks.append(('table', table_lines))
            continue
        if re.match(r'^\d+\.\s+', line):
            items = []
            while i < len(lines) and re.match(r'^\d+\.\s+', lines[i].strip()):
                items.append(re.sub(r'^\d+\.\s+', '', lines[i].strip()))
                i += 1
            blocks.append(('number', items))
            continue
        if re.match(r'^[-*]\s+', line):
            items = []
            while i < len(lines) and re.match(r'^[-*]\s+', lines[i].strip()):
                items.append(re.sub(r'^[-*]\s+', '', lines[i].strip()))
                i += 1
            blocks.append(('bullet', items))
            continue
        para = [line]
        i += 1
        while i < len(lines) and lines[i].strip() and not re.match(r'^(#{1,3}\s+|\|)|^```|^\d+\.\s+|^[-*]\s+', lines[i].strip()):
            para.append(lines[i].strip())
            i += 1
        blocks.append(('paragraph', ' '.join(para)))
    return blocks


def collect_toc(blocks: list[tuple[str, object]]) -> list[tuple[int, str]]:
    toc = []
    for kind, payload in blocks:
        if kind != 'heading':
            continue
        level, text, _raw = payload
        if text == TITLE_CN or text == '摘要' or text == 'Abstract':
            continue
        if text.startswith('第') or re.match(r'^\d+\.\d+', text):
            toc.append((level, text))
    return toc


def add_toc(doc: Document, toc_items: list[tuple[int, str]]) -> None:
    add_center_line(doc, '目  录', 16, True)
    for level, text in toc_items:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.left_indent = Cm(0.6 * max(level - 1, 0))
        r = p.add_run(text)
        set_run_font(r, east_asia='宋体', size=12)
    add_page_break(doc)


def add_paragraph_text(doc: Document, text: str, indent: bool = True, center: bool = False, bold: bool = False) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(0.74) if indent and not center else Cm(0)
    r = p.add_run(text)
    set_run_font(r, east_asia='宋体', size=12, bold=bold)


def add_image_with_caption(doc: Document, filename: str, caption: str) -> None:
    path = FIG_DIR / filename
    if not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Cm(14.5))
    cp = doc.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cp.add_run(caption)
    set_run_font(r, east_asia='宋体', size=10.5)


def add_table(doc: Document, table_lines: list[str]) -> None:
    rows = []
    for line in table_lines:
        stripped = line.strip('|')
        if set(stripped.replace('|', '').replace('-', '').replace(':', '').strip()) == set():
            continue
        rows.append([cell.strip() for cell in line.strip('|').split('|')])
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = 'Table Grid'
    for i, row in enumerate(rows):
        for j, cell in enumerate(row):
            p = table.cell(i, j).paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(cell.replace('`', ''))
            set_run_font(r, east_asia='宋体', size=10.5, bold=(i == 0))
    doc.add_paragraph()


def add_main_content(doc: Document, blocks: list[tuple[str, object]]) -> None:
    current_raw = ''
    for kind, payload in blocks:
        if kind == 'heading':
            level, text, raw = payload
            current_raw = raw
            if text == TITLE_CN:
                continue
            if text == '摘要':
                add_center_line(doc, '摘要', 16, True)
                continue
            if text == 'Abstract':
                add_page_break(doc)
                add_center_line(doc, 'Abstract', 16, True)
                continue
            style = {1: 'Heading 1', 2: 'Heading 2', 3: 'Heading 3'}.get(level, 'Heading 3')
            p = doc.add_paragraph(style=style)
            p.paragraph_format.first_line_indent = Cm(0)
            r = p.add_run(text)
            size = 16 if style == 'Heading 1' else 14 if style == 'Heading 2' else 12
            set_run_font(r, east_asia='黑体', size=size, bold=True)
            for filename, caption in IMAGE_PLAN.get(raw, []):
                add_image_with_caption(doc, filename, caption)
            continue
        if kind == 'paragraph':
            text = str(payload).replace('`', '')
            if text.startswith('**关键词**：'):
                add_paragraph_text(doc, text.replace('**', ''), indent=False)
            else:
                add_paragraph_text(doc, text)
            continue
        if kind == 'table':
            add_table(doc, payload)
            continue
        if kind in {'number', 'bullet'}:
            for idx, item in enumerate(payload):
                prefix = f'{idx + 1}. ' if kind == 'number' else '• '
                add_paragraph_text(doc, prefix + str(item).replace('`', ''), indent=False)
            continue
        if kind == 'code':
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Cm(0)
            r = p.add_run(str(payload))
            set_run_font(r, east_asia='宋体', latin='Courier New', size=10.5)


def main() -> None:
    template = Document(TEMPLATE_PATH)
    doc = Document()
    clear_document(doc)
    configure_styles(doc, template)

    add_cover(doc)

    md = read_and_refine_markdown()
    blocks = parse_blocks(md)
    toc_items = collect_toc(blocks)
    add_toc(doc, toc_items)
    add_main_content(doc, blocks)

    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT_PATH)
    print(f'Saved standardized thesis to: {OUTPUT_PATH}')


if __name__ == '__main__':
    main()
