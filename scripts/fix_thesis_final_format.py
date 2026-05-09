from __future__ import annotations

import re
import shutil
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.table import Table
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[1]
BACKUP_PATH = ROOT / "output" / "doc" / "thesis_backup_before_final_format_fix.docx"
OUTPUT_COPY = ROOT / "output" / "doc" / "thesis_final_format_fixed.docx"

THESIS_TITLE = "基于深度学习的吸烟者行为检测系统设计"
EVEN_HEADER = "莆田学院本科生毕业设计说明书"

TABLE_CAPTIONS = {
    "表3-1": "表3-1 主要原始数据来源及作用定位",
    "表3-2": "表3-2 主要数据源类别映射规则",
    "表3-3": "表3-3 数据清洗阶段异常情况",
    "表3-4": "表3-4 全量数据集与温和均衡版数据集统计对比",
    "表3-5": "表3-5 smoke_bal数据集划分与类别统计",
    "表3-6": "表3-6 Baseline训练参数设置",
    "表4-1": "表4-1 Baseline与ECA整体指标对比",
    "表4-2": "表4-2 Baseline与ECA分类别指标对比",
    "表4-3": "表4-3 测试集cigarette目标尺寸分布",
    "表5-1": "表5-1 系统主要数据表及作用",
    "表5-2": "表5-2 系统测试项目及结果",
}

FIGURE_CAPTIONS = {
    "图2-1": "图2-1 YOLOv8各版本性能对比",
    "图2-2": "图2-2 YOLOv8n网络结构",
    "图2-3": "图2-3 ECA-Net模块结构",
    "图4-1": "图4-1 Baseline模型测试集归一化混淆矩阵",
    "图4-2": "图4-2 Baseline模型训练结果曲线",
    "图5-1": "图5-1 网页端图片检测结果示意图",
    "图5-2": "图5-2 系统总体架构图",
    "图5-3": "图5-3 数据库关系说明图",
}


def find_target_docx() -> Path:
    matches = [
        path
        for path in ROOT.glob("*.docx")
        if "224-" in path.name and not path.name.startswith("~$") and path.stat().st_size > 1000
    ]
    if not matches:
        raise FileNotFoundError("未找到当前定稿论文 DOCX 文件。")
    matches.sort()
    return matches[0]


def set_run_font(run, east_asia: str = "宋体", latin: str = "Times New Roman", size: float = 10.5, bold: bool | None = None) -> None:
    run.font.name = latin
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), east_asia)
    rfonts.set(qn("w:ascii"), latin)
    rfonts.set(qn("w:hAnsi"), latin)


def apply_para_font(paragraph: Paragraph, east_asia: str = "宋体", latin: str = "Times New Roman", size: float = 10.5, bold: bool | None = None) -> None:
    if not paragraph.runs:
        paragraph.add_run("")
    for run in paragraph.runs:
        set_run_font(run, east_asia=east_asia, latin=latin, size=size, bold=bold)


def set_para_format(
    paragraph: Paragraph,
    *,
    alignment: WD_ALIGN_PARAGRAPH | None = None,
    first_indent: float | None = None,
    line_spacing: float | None = 1.5,
    before: float = 0,
    after: float = 0,
) -> None:
    if alignment is not None:
        paragraph.alignment = alignment
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line_spacing
    fmt.first_line_indent = Pt(first_indent) if first_indent is not None else None


def reset_part_to_one_paragraph(part) -> Paragraph:
    element = part._element
    for child in list(element):
        element.remove(child)
    p = OxmlElement("w:p")
    element.append(p)
    return Paragraph(p, part)


def add_page_field(paragraph: Paragraph) -> None:
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)
    set_run_font(run, size=10.5)


def enable_odd_even_headers(doc: Document) -> None:
    settings = doc.settings.element
    if settings.find(qn("w:evenAndOddHeaders")) is None:
        settings.append(OxmlElement("w:evenAndOddHeaders"))
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")


def format_headers_footers(doc: Document) -> None:
    enable_odd_even_headers(doc)
    for section in doc.sections:
        section.different_first_page_header_footer = False
        section.header.is_linked_to_previous = False
        section.even_page_header.is_linked_to_previous = False
        section.footer.is_linked_to_previous = False
        section.even_page_footer.is_linked_to_previous = False
        section.first_page_footer.is_linked_to_previous = False

        for part, text in [(section.header, THESIS_TITLE), (section.even_page_header, EVEN_HEADER)]:
            p = reset_part_to_one_paragraph(part)
            p.text = text
            set_para_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.0)
            apply_para_font(p, east_asia="宋体", size=9)

        for part in [section.footer, section.even_page_footer]:
            p = reset_part_to_one_paragraph(part)
            set_para_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.0)
            add_page_field(p)

        reset_part_to_one_paragraph(section.first_page_footer)


def format_styles_and_paragraphs(doc: Document) -> None:
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        style = paragraph.style.name

        if not text:
            set_para_format(paragraph, line_spacing=1.0, before=0, after=0)
            continue

        if style == "Heading 1" or text in {"参考文献", "致谢"}:
            paragraph.style = "Heading 1"
            set_para_format(paragraph, alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.5, before=12, after=6)
            apply_para_font(paragraph, east_asia="黑体", size=14, bold=False)
            continue

        if style == "Heading 2":
            set_para_format(paragraph, alignment=WD_ALIGN_PARAGRAPH.LEFT, line_spacing=1.5, before=6, after=0)
            apply_para_font(paragraph, east_asia="黑体", size=12, bold=False)
            continue

        if re.match(r"^\[\d+\]", text):
            set_para_format(paragraph, alignment=WD_ALIGN_PARAGRAPH.LEFT, first_indent=None, line_spacing=1.25, before=0, after=0)
            paragraph.paragraph_format.left_indent = Pt(18)
            paragraph.paragraph_format.first_line_indent = Pt(-18)
            apply_para_font(paragraph, east_asia="宋体", size=10.5)
            continue

        if is_table_caption(text):
            set_para_format(paragraph, alignment=WD_ALIGN_PARAGRAPH.CENTER, first_indent=None, line_spacing=1.25, before=6, after=3)
            apply_para_font(paragraph, east_asia="宋体", size=10.5, bold=False)
            continue

        if is_figure_caption(text):
            set_para_format(paragraph, alignment=WD_ALIGN_PARAGRAPH.CENTER, first_indent=None, line_spacing=1.25, before=3, after=6)
            apply_para_font(paragraph, east_asia="宋体", size=10.5, bold=False)
            continue

        if re.search(r"（4-\d+）$", text):
            set_para_format(paragraph, alignment=WD_ALIGN_PARAGRAPH.CENTER, first_indent=None, line_spacing=1.25, before=3, after=3)
            apply_para_font(paragraph, east_asia="宋体", size=10.5)
            continue

        set_para_format(paragraph, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, first_indent=21, line_spacing=1.5, before=0, after=0)
        apply_para_font(paragraph, east_asia="宋体", size=10.5)


def format_front_matter(doc: Document) -> None:
    for idx, paragraph in enumerate(doc.paragraphs[:55]):
        text = paragraph.text.strip()
        if not text:
            continue

        if text == "本 科 生 毕 业 设 计 说 明 书":
            set_para_format(paragraph, alignment=WD_ALIGN_PARAGRAPH.CENTER, first_indent=None, line_spacing=1.5, before=0, after=0)
            apply_para_font(paragraph, east_asia="宋体", size=26, bold=True)
            continue

        if idx == 4:
            set_para_format(paragraph, alignment=WD_ALIGN_PARAGRAPH.CENTER, first_indent=None, line_spacing=1.5, before=0, after=0)
            apply_para_font(paragraph, east_asia="楷体", size=18, bold=True)
            continue

        if text == "陈刚" and idx < 10:
            set_para_format(paragraph, alignment=WD_ALIGN_PARAGRAPH.CENTER, first_indent=None, line_spacing=1.5, before=0, after=0)
            apply_para_font(paragraph, east_asia="楷体", size=18, bold=True)
            continue

        if text.startswith("学    号："):
            paragraph.text = "学    号：     202212903403"
            text = paragraph.text.strip()

        if any(text.startswith(prefix) for prefix in ["学    院：", "专    业：", "班    级：", "学    号：", "指导教师：", "职    称："]):
            set_para_format(paragraph, alignment=WD_ALIGN_PARAGRAPH.LEFT, first_indent=72, line_spacing=1.5, before=0, after=0)
            apply_para_font(paragraph, east_asia="楷体", size=16, bold=True)
            continue

        if re.match(r"^20\d{2}年\s*\d+\s*月$", text):
            set_para_format(paragraph, alignment=WD_ALIGN_PARAGRAPH.CENTER, first_indent=None, line_spacing=1.5, before=0, after=0)
            apply_para_font(paragraph, east_asia="楷体", size=16, bold=True)
            continue

        if text in {"原创性声明", "指导声明"}:
            set_para_format(paragraph, alignment=WD_ALIGN_PARAGRAPH.CENTER, first_indent=None, line_spacing=1.5, before=6, after=6)
            apply_para_font(paragraph, east_asia="黑体", size=18, bold=False)
            continue

        if text.startswith("本人郑重声明") or text.startswith("本人指导的"):
            set_para_format(paragraph, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, first_indent=32, line_spacing=1.5, before=0, after=0)
            apply_para_font(paragraph, east_asia="仿宋", size=14, bold=False)
            continue

        if text.startswith("学生签名") or text.startswith("指导教师签名"):
            set_para_format(paragraph, alignment=WD_ALIGN_PARAGRAPH.CENTER, first_indent=None, line_spacing=1.5, before=0, after=0)
            apply_para_font(paragraph, east_asia="仿宋", size=14, bold=False)
            continue

        if idx in {38, 45}:
            set_para_format(paragraph, alignment=WD_ALIGN_PARAGRAPH.CENTER, first_indent=None, line_spacing=1.5, before=6, after=3)
            apply_para_font(paragraph, east_asia="黑体", latin="Times New Roman", size=14, bold=True)
            continue

        if idx in {39, 40, 46, 47}:
            set_para_format(paragraph, alignment=WD_ALIGN_PARAGRAPH.CENTER, first_indent=None, line_spacing=1.5, before=0, after=0)
            apply_para_font(paragraph, east_asia="宋体", latin="Times New Roman", size=10.5, bold=False)
            continue

        if idx in {42, 48, 49, 50}:
            set_para_format(paragraph, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, first_indent=21, line_spacing=1.5, before=0, after=0)
            apply_para_font(paragraph, east_asia="宋体", latin="Times New Roman", size=10.5, bold=False)
            continue

        if idx in {43, 53}:
            set_para_format(paragraph, alignment=WD_ALIGN_PARAGRAPH.LEFT, first_indent=None, line_spacing=1.5, before=3, after=0)
            apply_para_font(paragraph, east_asia="宋体", latin="Times New Roman", size=10.5, bold=False)
            continue


def is_table_caption(text: str) -> bool:
    return re.match(r"^表\d+-\d+\s+", text) is not None


def is_figure_caption(text: str) -> bool:
    return re.match(r"^图\d+-\d+\s+", text) is not None


def normalize_captions(doc: Document) -> None:
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        for key, replacement in TABLE_CAPTIONS.items():
            if text.startswith(key):
                paragraph.text = replacement
                break
        for key, replacement in FIGURE_CAPTIONS.items():
            if text.startswith(key + " ") or text == key:
                paragraph.text = replacement
                break
    dedupe_figure_captions(doc)


def dedupe_figure_captions(doc: Document) -> None:
    seen: dict[str, int] = {}
    for idx, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        match = re.match(r"^(图\d+-\d+)\s+(.+)$", text)
        if not match:
            continue
        key = match.group(1)
        if key not in seen:
            seen[key] = idx
            continue
        if key == "图5-1":
            doc.paragraphs[seen[key]].text = (
                "网页端图片检测结果如图5-1所示。上传后的图像会在页面中返回标注结果，"
                "同时保留记录编号和文件路径，便于在答辩时说明样例来自哪一次检测。"
                "这样既能展示视觉效果，也能展示系统对结果的管理方式。"
            )
            seen[key] = idx


def add_citations(doc: Document) -> None:
    edits = [
        ("两阶段方法如 Faster R-CNN", "两阶段方法如 Faster R-CNN[5]"),
        ("单阶段方法如 SSD 和 YOLO 系列", "单阶段方法如 SSD[6] 和 YOLO 系列[1-4]"),
        ("在本课题中，卷积神经网络主要承担特征提取任务。", "在本课题中，卷积神经网络主要承担特征提取任务[8-10]。"),
        ("YOLO 系列属于单阶段检测框架", "YOLO 系列属于单阶段检测框架[1-4]"),
        ("ECA 属于通道注意力方法", "ECA 属于通道注意力方法[7]"),
        ("可以继续尝试更高输入分辨率、多尺度训练、知识蒸馏或切片辅助检测等方法。", "可以继续尝试更高输入分辨率、多尺度训练、知识蒸馏[11-13]或切片辅助检测[14]等方法。"),
    ]
    for paragraph in doc.paragraphs:
        text = paragraph.text
        for old, new in edits:
            if old in text and new not in text:
                paragraph.text = text.replace(old, new)
                text = paragraph.text
        if "YOLO 系列[1-4]则" in text:
            paragraph.text = text.replace("YOLO 系列[1-4]则", "YOLO 系列[1-4]，则")
            text = paragraph.text


def insert_after(paragraph: Paragraph, text: str, style: str = "Normal") -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    new_para.style = style
    new_para.add_run(text)
    return new_para


def insert_formula_block(doc: Document) -> None:
    if any("Precision = TP" in p.text for p in doc.paragraphs):
        return
    target = next((p for p in doc.paragraphs if "本文使用 Precision、Recall、mAP@0.5" in p.text), None)
    if target is None:
        return
    target.text = target.text.rstrip("。") + "，其基本计算形式见式（4-1）至式（4-3）。"
    current = target
    for formula in [
        "Precision = TP / (TP + FP)        （4-1）",
        "Recall = TP / (TP + FN)        （4-2）",
        "mAP = (1 / N) Σ AP_i        （4-3）",
    ]:
        current = insert_after(current, formula)


def ensure_acknowledgement(doc: Document) -> None:
    if any(p.text.strip() == "致谢" for p in doc.paragraphs):
        return
    ref = next((p for p in doc.paragraphs if p.text.strip() == "参考文献"), None)
    if ref is None:
        return
    ack_body = (
        "在本次毕业设计完成过程中，指导教师林一锋老师在选题确定、技术路线、实验分析和论文修改等方面给予了耐心指导。"
        "同时，学院老师和同学在数据整理、环境配置和系统测试过程中也提供了帮助。"
        "在此向所有关心和帮助过本课题完成的老师、同学表示感谢。"
    )
    ack_heading = ref.insert_paragraph_before("致谢")
    ack_heading.style = "Heading 1"
    ack_heading.paragraph_format.page_break_before = True
    ack_para = insert_after(ack_heading, ack_body)
    ack_para.style = "Normal"
    ref.paragraph_format.page_break_before = True


def format_references(doc: Document) -> None:
    replacements = {
        1: "[1] Redmon J, Divvala S, Girshick R, et al. You Only Look Once: Unified, Real-Time Object Detection[C]//Proceedings of the IEEE Conference on Computer Vision and Pattern Recognition. Las Vegas: IEEE, 2016: 779-788.",
        2: "[2] Redmon J, Farhadi A. YOLOv3: An Incremental Improvement[EB/OL]. (2018-04-08)[2026-04-27]. https://arxiv.org/abs/1804.02767.",
        3: "[3] Bochkovskiy A, Wang C Y, Liao H Y M. YOLOv4: Optimal Speed and Accuracy of Object Detection[EB/OL]. (2020-04-23)[2026-04-27]. https://arxiv.org/abs/2004.10934.",
        4: "[4] Jocher G, Chaurasia A, Qiu J. Ultralytics YOLOv8[EB/OL]. (2023-01-10)[2026-04-27]. https://github.com/ultralytics/ultralytics.",
        5: "[5] Ren S, He K, Girshick R, et al. Faster R-CNN: Towards Real-Time Object Detection with Region Proposal Networks[J]. IEEE Transactions on Pattern Analysis and Machine Intelligence, 2017, 39(6): 1137-1149.",
        6: "[6] Liu W, Anguelov D, Erhan D, et al. SSD: Single Shot MultiBox Detector[C]//Computer Vision - ECCV 2016. Cham: Springer, 2016: 21-37.",
        7: "[7] Wang Q, Wu B, Zhu P, et al. ECA-Net: Efficient Channel Attention for Deep Convolutional Neural Networks[C]//Proceedings of the IEEE/CVF Conference on Computer Vision and Pattern Recognition. Seattle: IEEE, 2020: 11531-11539.",
        8: "[8] LeCun Y, Bottou L, Bengio Y, et al. Gradient-Based Learning Applied to Document Recognition[J]. Proceedings of the IEEE, 1998, 86(11): 2278-2324.",
        9: "[9] Krizhevsky A, Sutskever I, Hinton G E. ImageNet Classification with Deep Convolutional Neural Networks[J]. Communications of the ACM, 2017, 60(6): 84-90.",
        10: "[10] He K, Zhang X, Ren S, et al. Deep Residual Learning for Image Recognition[C]//Proceedings of the IEEE Conference on Computer Vision and Pattern Recognition. Las Vegas: IEEE, 2016: 770-778.",
        11: "[11] Hinton G, Vinyals O, Dean J. Distilling the Knowledge in a Neural Network[EB/OL]. (2015-03-09)[2026-04-27]. https://arxiv.org/abs/1503.02531.",
        12: "[12] Yang Z, Li Z, Jiang X, et al. Focal and Global Knowledge Distillation for Detectors[C]//Proceedings of the IEEE/CVF Conference on Computer Vision and Pattern Recognition. New Orleans: IEEE, 2022: 4643-4652.",
        13: "[13] Xu M, Zhang Z, Hu H, et al. End-to-End Semi-Supervised Object Detection With Soft Teacher[C]//Proceedings of the IEEE/CVF International Conference on Computer Vision. Montreal: IEEE, 2021: 3060-3069.",
        14: "[14] Akyon F C, Altinuc S O, Temizel A. Slicing Aided Hyper Inference and Fine-Tuning for Small Object Detection[C]//2022 IEEE International Conference on Image Processing. Bordeaux: IEEE, 2022: 966-970.",
    }
    ref_re = re.compile(r"^\[(\d+)\]")
    for paragraph in doc.paragraphs:
        match = ref_re.match(paragraph.text.strip())
        if match:
            number = int(match.group(1))
            if number in replacements:
                paragraph.text = replacements[number]


def set_cell_border(cell, **kwargs) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        if edge not in kwargs:
            continue
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        for key, value in kwargs[edge].items():
            element.set(qn("w:" + key), str(value))


def format_three_line_tables(doc: Document) -> None:
    none = {"val": "nil"}
    top = {"val": "single", "sz": "12", "space": "0", "color": "000000"}
    mid = {"val": "single", "sz": "8", "space": "0", "color": "000000"}
    bottom = {"val": "single", "sz": "12", "space": "0", "color": "000000"}
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                set_cell_border(cell, top=none, left=none, bottom=none, right=none, insideH=none, insideV=none)
                for p in cell.paragraphs:
                    set_para_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER, first_indent=None, line_spacing=1.25, before=0, after=0)
                    apply_para_font(p, east_asia="宋体", size=10.5)
        if not table.rows:
            continue
        for cell in table.rows[0].cells:
            set_cell_border(cell, top=top, bottom=mid, left=none, right=none)
            for p in cell.paragraphs:
                apply_para_font(p, east_asia="宋体", size=10.5, bold=True)
        for cell in table.rows[-1].cells:
            set_cell_border(cell, bottom=bottom, left=none, right=none)


def set_section_layout(doc: Document) -> None:
    for section in doc.sections:
        section.start_type = WD_SECTION.NEW_PAGE
        section.top_margin = Pt(56.7)
        section.bottom_margin = Pt(56.7)
        section.left_margin = Pt(70.9)
        section.right_margin = Pt(56.7)


def fix_doc(doc_path: Path) -> None:
    doc = Document(doc_path)
    set_section_layout(doc)
    normalize_captions(doc)
    add_citations(doc)
    insert_formula_block(doc)
    ensure_acknowledgement(doc)
    format_references(doc)
    format_headers_footers(doc)
    format_three_line_tables(doc)
    format_styles_and_paragraphs(doc)
    format_front_matter(doc)
    doc.save(doc_path)


def main() -> None:
    target = find_target_docx()
    BACKUP_PATH.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(target, BACKUP_PATH)
    OUTPUT_COPY.parent.mkdir(parents=True, exist_ok=True)
    try:
        fix_doc(target)
        shutil.copy2(target, OUTPUT_COPY)
        saved_target = target
    except PermissionError:
        temp = OUTPUT_COPY
        shutil.copy2(target, temp)
        fix_doc(temp)
        saved_target = temp
    print(f"Backup saved to: {BACKUP_PATH}")
    print(f"Fixed thesis saved to: {saved_target}")
    print(f"Output copy saved to: {OUTPUT_COPY}")


if __name__ == "__main__":
    main()
