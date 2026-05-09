from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[1]

ABSTRACT_CN_FINAL = (
    "在医院、校园、办公楼和交通站点等禁烟场景中，吸烟行为管理一直比较依赖人工巡查。"
    "这样的方式虽然直观，但在实际使用中往往存在覆盖范围有限、发现不够及时和长期维护成本较高等问题。"
    "因此，利用视频监控画面自动识别吸烟行为，具有一定的现实意义。结合本次毕业设计任务，"
    "本文围绕吸烟行为检测开展了数据整理、模型训练和展示方案设计等工作。"
    "本文以 YOLOv8n 为基础模型，将检测目标统一为 cigarette、smoking_person 和 smoke 三类。"
    "针对项目中不同来源数据格式不一致、类别划分不统一的问题，完成了数据转换、标签映射、异常标签清理、"
    "数据合并和数据集划分，构建了可直接用于训练的吸烟行为检测数据集。考虑到本地实验环境以 CPU 为主，"
    "训练耗时较长，本文又在全量数据集的基础上构建了温和均衡版数据集 smoke_bal，用于当前阶段的实验。"
    "在模型部分，本文先搭建了 YOLOv8n 的训练、验证、推理和导出流程，再在此基础上引入 ECA 通道注意力模块，"
    "设计 YOLOv8n+ECA 改进模型，用于对比实验。当前已经完成 baseline 模型和 ECA 模型在 smoke_bal 数据集上的 "
    "30 轮训练与测试验证。结果表明，baseline 模型的 Precision 为 0.526，Recall 为 0.625，mAP@0.5 为 0.520，"
    "mAP@0.5:0.95 为 0.323；ECA 模型的 Precision 为 0.513，Recall 为 0.572，mAP@0.5 为 0.494，"
    "mAP@0.5:0.95 为 0.299。整体来看，当前 baseline 模型表现更稳定，smoke 类效果较好，而 cigarette 类仍然是"
    "本课题中的主要难点。在系统实现方面，项目已经完成训练脚本、验证脚本、预测脚本和模型导出脚本的搭建，并形成了"
    "基于 FastAPI 的网页展示方案，用于展示实验流程、检测结果和案例演示。总体来看，本文已经完成了毕业设计的主要框架，"
    "能够为论文定稿和后续答辩准备提供支撑。"
)

ABSTRACT_EN_FINAL = (
    "In system implementation, the project has completed the training, validation, prediction and export scripts, "
    "and adopts a FastAPI-based web interface as the final presentation layer for thesis defense, which is more suitable "
    "for demonstration and explanation."
)

KEYWORDS_CN_FINAL = "关键词：吸烟行为检测；YOLOv8n；ECA；目标检测；FastAPI"
MAIN_WORK_ITEM_5 = "5. 结合答辩展示需求，设计并实现基于 FastAPI 的网页展示方案，用于统一呈现实验流程、检测结果与系统状态。"
REFERENCES_TITLE = "参考文献"
APPENDIX_TITLE = "附录：项目对应的关键命令"


def find_target() -> Path:
    matches = [p for p in ROOT.glob("定稿_智能224-*.docx") if not p.name.startswith("~$")]
    if not matches:
        raise FileNotFoundError("No final docx found.")
    return matches[0]


def insert_after(paragraph, text: str, style: str | None = None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = paragraph._parent.add_paragraph()
    new_para._p = new_p
    if text:
        new_para.add_run(text)
    if style:
        new_para.style = style
    return new_para


def set_update_fields(doc: Document) -> None:
    settings = doc.settings.element
    found = settings.find(qn("w:updateFields"))
    if found is None:
        upd = OxmlElement("w:updateFields")
        upd.set(qn("w:val"), "true")
        settings.append(upd)
    else:
        found.set(qn("w:val"), "true")


def normalize_cover(doc: Document) -> None:
    replacements = {
        "基于深度学习的吸烟者行为检测": "基于深度学习的吸烟者行为检测系统设计与实现",
        "关键词：深度学习；遥感检测；YOLOv8s；迭代注意力机制": "关键词：吸烟行为检测；YOLOv8n；ECA；目标检测；FastAPI",
        "Keywords:  Deep Learning; Remote Sensing Detection; YOLOv8s; Iterative Attention Mechanism": "Keywords: Deep Learning; Smoking Behavior Detection; YOLOv8n; ECA; FastAPI",
    }
    for p in doc.paragraphs:
        t = p.text.strip()
        if t in replacements:
            p.text = replacements[t]
        elif t.startswith("摘要：吸烟行为对公共健康造成严重危害"):
            p.text = (
                "摘要：在医院、校园、办公楼和交通站点等禁烟场景中，吸烟行为管理长期依赖人工巡查，"
                "存在覆盖范围有限、发现不够及时和持续成本较高等问题。本文围绕吸烟行为检测任务，"
                "基于 YOLOv8n 搭建训练、验证、推理和导出流程，并对多来源数据完成转换、标签映射、"
                "异常标签清理和统一划分。在此基础上构建温和均衡版数据集 smoke_bal，完成 baseline 与 "
                "ECA 改进模型的 30 轮训练和测试验证。实验结果表明，当前 baseline 模型整体表现更稳定，"
                "smoke 类检测效果较好，而 cigarette 类仍然是主要难点。在系统展示方面，本文形成了基于 "
                "FastAPI 的网页展示方案，用于统一呈现实验流程、检测结果和案例演示。"
            )


def repair_front_matter(doc: Document) -> None:
    for p in doc.paragraphs:
        t = p.text.strip()
        if not t:
            continue

        if ("YOLOv8n" in t and "smoke_bal" in t and ("?" in t or "PyQt5" in t)) or t.startswith("摘要："):
            p.text = ABSTRACT_CN_FINAL
            continue

        if t == ABSTRACT_EN_FINAL:
            p.text = ABSTRACT_EN_FINAL
            continue

        if t.startswith("关键词：") or t.startswith("**关键词**："):
            p.text = KEYWORDS_CN_FINAL
            continue

        if (t.startswith("5.") and "FastAPI" in t) or ("PyQt5" in t and t.startswith("5.")) or ("?" in t and t.startswith("5.")):
            p.text = MAIN_WORK_ITEM_5
            continue


def set_run_east_asia(run, east_asia: str = "宋体") -> None:
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), east_asia)


def insert_paragraph_after(paragraph, text: str = ""):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = paragraph._parent.add_paragraph()
    new_para._p = new_p
    if text:
        new_para.add_run(text)
    return new_para


def normalize_reference_entry_text(text: str) -> str:
    cleaned = text.replace("\r", "\n")
    cleaned = cleaned.strip()
    if cleaned.startswith("[") and "]" in cleaned:
        cleaned = cleaned.split("]", 1)[1].strip()
    if cleaned.startswith("［") and "］" in cleaned:
        cleaned = cleaned.split("］", 1)[1].strip()
    return cleaned


def normalize_references(doc: Document) -> None:
    heading_idx = None
    appendix_idx = None
    for idx, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text == REFERENCES_TITLE and heading_idx is None:
            heading_idx = idx
        elif text == APPENDIX_TITLE and appendix_idx is None:
            appendix_idx = idx
            break

    if heading_idx is None:
        return

    heading_para = doc.paragraphs[heading_idx]
    heading_para.style = "Normal"
    heading_para.alignment = None
    heading_fmt = heading_para.paragraph_format
    heading_fmt.left_indent = None
    heading_fmt.right_indent = None
    heading_fmt.first_line_indent = None
    heading_fmt.line_spacing = None
    heading_fmt.space_before = None
    heading_fmt.space_after = None
    if heading_para.runs:
        for run in heading_para.runs:
            run.bold = False
            run.font.size = Pt(12)
            run.font.name = "Times New Roman"
            set_run_east_asia(run, "宋体")

    entry_end = appendix_idx if appendix_idx is not None else len(doc.paragraphs)
    raw_lines: list[str] = []
    for para in doc.paragraphs[heading_idx + 1:entry_end]:
        text = para.text.replace("\u000b", "\n")
        parts = []
        for line in text.splitlines():
            if line.strip():
                parts.append(line.strip())
        if not parts:
            continue
        merged = " ".join(parts)
        chunks = re.split(r"(?=(?:\[\d+\]|［\d+］)\s*)", merged)
        for chunk in chunks:
            cleaned = normalize_reference_entry_text(chunk)
            if cleaned:
                raw_lines.append(cleaned)

    if not raw_lines:
        return

    # Reuse the first reference paragraph if present, clear the rest.
    first_entry_para = doc.paragraphs[heading_idx + 1] if heading_idx + 1 < len(doc.paragraphs) else insert_paragraph_after(heading_para)
    for para in doc.paragraphs[heading_idx + 1:entry_end]:
        para.text = ""
        para.style = "Normal"

    current = first_entry_para
    for idx, entry in enumerate(raw_lines):
        if idx == 0:
            current.text = entry
        else:
            current = insert_paragraph_after(current, entry)
        current.style = "Normal"
        current.alignment = None
        fmt = current.paragraph_format
        fmt.left_indent = None
        fmt.right_indent = None
        fmt.first_line_indent = None
        fmt.line_spacing = None
        fmt.space_before = None
        fmt.space_after = None
        for run in current.runs:
            run.bold = False
            run.font.size = Pt(9)
            run.font.name = "Times New Roman"
            set_run_east_asia(run, "宋体")


def normalize_headings(doc: Document) -> None:
    mapping = {
        "1 引言": ("第1章 绪论", "Heading 1"),
        "1.1 研究背景": ("1.1 研究背景", "Heading 2"),
        "1.2 国内外应用现状": ("1.2 国内外研究现状", "Heading 2"),
        "1.3 主要工作和论文结构": ("1.3 本文主要工作", "Heading 2"),
        "2 相关理论和技术": ("第2章 相关理论与关键技术", "Heading 1"),
        "2.1 卷积神经网络与目标检测基础": ("2.1 卷积神经网络与目标检测基础", "Heading 2"),
        "2.2 YOLO目标检测算法": ("2.2 YOLOv8n 模型原理", "Heading 2"),
        "2.4 ECA注意力机制": ("2.3 ECA 注意力机制", "Heading 2"),
        "2.5 FastAPI 与 ONNX 展示部署技术": ("2.4 FastAPI 与 ONNX 展示部署技术", "Heading 2"),
        "3 数据集的构建和改进模型设计": ("第3章 数据集构建与改进模型设计", "Heading 1"),
        "3.1 数据来源与类别统一": ("3.1 数据来源与类别统一", "Heading 2"),
        "3.2 数据处理流程设计": ("3.2 数据处理流程设计", "Heading 2"),
        "3.4 最终数据集统计分析": ("3.3 最终数据集统计分析", "Heading 2"),
        "3.4 YOLOv8n baseline 训练框架设计": ("3.4 YOLOv8n baseline 训练框架设计", "Heading 2"),
        "3.5YOLOv8n+ECA 改进模型设计": ("3.5 YOLOv8n+ECA 改进模型设计", "Heading 2"),
        "4 实验设计和结果分析": ("第4章 实验设计与结果分析", "Heading 1"),
        "4.1实验设备及环境": ("4.1 实验环境", "Heading 2"),
        "4.2评价指标": ("4.2 评价指标", "Heading 2"),
        "4.3 baseline 与 ECA 实验结果": ("4.3 baseline 与 ECA 实验结果", "Heading 2"),
        "4.4 结果分析": ("4.4 结果分析", "Heading 2"),
        "4.5 本章小结": ("4.5 本章小结", "Heading 2"),
        "5系统实现与展现设计": ("第5章 系统实现与展示设计", "Heading 1"),
        "5.1 系统设计": ("5.1 系统总体架构", "Heading 2"),
        "5.2推理与导出流程实现": ("5.2 推理与导出流程实现", "Heading 2"),
        "5.3 PyQt5 原型与网页展示方案": ("5.3 PyQt5 原型与网页展示方案", "Heading 2"),
        "5.4 FastAPI 前端展示设计": ("5.4 FastAPI 前端展示设计", "Heading 2"),
        "5.5 本章小结": ("5.5 本章小结", "Heading 2"),
        "第6章 总结与展望": ("第6章 总结与展望", "Heading 1"),
        "6.1 全文总结": ("6.1 全文总结", "Heading 2"),
        "6.2 不足与展望": ("6.2 不足与展望", "Heading 2"),
    }
    for p in doc.paragraphs:
        t = p.text.strip()
        if t in mapping:
            text, style = mapping[t]
            p.text = text
            p.style = style
        elif t == "1.3.1 主要工作":
            p.text = ""
            p.style = "Normal"


def fix_mixed_outline(doc: Document) -> None:
    for p in doc.paragraphs:
        t = p.text.strip()
        if "1.4 论文结构安排" not in t:
            continue
        parts = t.split("1.4 论文结构安排", 1)
        before = parts[0].strip()
        after = parts[1].strip()
        if before:
            p.text = before
            p.style = "Normal"
        else:
            p.text = ""
            p.style = "Normal"
        p2 = insert_after(p, "1.4 论文结构安排", "Heading 2")
        insert_after(p2, after, "Normal")
        break


def refine_54(doc: Document) -> None:
    for idx, p in enumerate(doc.paragraphs):
        if p.text.strip() != "5.4 FastAPI 前端展示设计":
            continue
        p1 = doc.paragraphs[idx + 1]
        p2 = doc.paragraphs[idx + 2]
        p1.text = (
            "从答辩展示需求出发，FastAPI 前端展示并不是一个面向复杂业务场景的完整系统，"
            "而是一个更适合毕业设计演示的轻量网页界面。后端负责加载训练好的权重文件，"
            "接收图片或视频请求并调用现有推理流程，前端负责展示原图、检测结果图、类别信息和关键指标。"
            "这样既能够最大限度复用现有 Python 代码，也便于在本机环境中快速部署和稳定演示。"
        )
        p2.text = (
            "从页面结构上看，展示页可以划分为课题简介、系统流程、检测演示和实验结果四个模块。"
            "课题简介页面用于说明研究背景和研究目标；系统流程页面展示数据整理、模型训练和测试验证的整体过程；"
            "检测演示页面支持上传图片并返回带检测框的结果图；实验结果页面用于展示 baseline 与 ECA 的对比指标、"
            "典型检测样例以及误检案例。按照这样的顺序组织页面，更符合答辩时“为什么做、怎么做、效果如何”的讲解逻辑。"
        )
        p3 = doc.paragraphs[idx + 3]
        if p3.text.strip().startswith("5.5 "):
            insert_after(
                p2,
                "在部署方式上，该方案不依赖复杂数据库和分布式服务，只需在本机完成模型加载、接口启动和页面访问即可，"
                "比较适合毕业答辩这种时间较短、稳定性优先的展示场景。后续如果还有时间，可以继续补充视频检测、结果导出和历史记录等功能，"
                "但现阶段应以保证演示流程完整、操作简洁和结果可展示为主。",
                "Normal",
            )
        else:
            p3.text = (
                "在部署方式上，该方案不依赖复杂数据库和分布式服务，只需在本机完成模型加载、接口启动和页面访问即可，"
                "比较适合毕业答辩这种时间较短、稳定性优先的展示场景。后续如果还有时间，可以继续补充视频检测、结果导出和历史记录等功能，"
                "但现阶段应以保证演示流程完整、操作简洁和结果可展示为主。"
            )
        break


def main() -> None:
    target = find_target()
    doc = Document(target)
    set_update_fields(doc)
    normalize_cover(doc)
    repair_front_matter(doc)
    normalize_headings(doc)
    fix_mixed_outline(doc)
    refine_54(doc)
    normalize_references(doc)
    doc.save(target)
    print(target)


if __name__ == "__main__":
    main()
