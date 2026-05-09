from __future__ import annotations

import argparse
import re
import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
TARGET_GLOB = "定稿_智能224-*.docx"
DEFAULT_OUTPUT = ROOT / "output" / "doc" / "thesis_sample_aligned.docx"
BACKUP_PATH = ROOT / "output" / "doc" / "thesis_backup_before_sample_format_alignment.docx"

CHAPTER_RE = re.compile(r"^第\s*(\d+)\s*章\s*(.+)$")
NORMALIZED_CHAPTER_RE = re.compile(r"^(\d+)\s+(.+)$")
SECTION_RE = re.compile(r"^\d+\.\d+\s+.+$")
SUBSECTION_RE = re.compile(r"^\d+\.\d+\.\d+\s+.+$")
CAPTION_RE = re.compile(r"^[图表]\s*\d+(?:[-－—]\d+)?\s+.+$")
APPENDIX_RE = re.compile(r"^附录")

REFERENCES_TITLE = "参考文献"
EN_TITLE_PREFIX = "Design of "


def find_target_docx() -> Path:
    matches = [p for p in ROOT.glob(TARGET_GLOB) if not p.name.startswith("~$")]
    if not matches:
        raise FileNotFoundError("未找到当前论文定稿 DOCX 文件。")
    return matches[0]


def backup_docx(target: Path) -> Path:
    BACKUP_PATH.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(target, BACKUP_PATH)
    return BACKUP_PATH


def ensure_output_parent(path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)


def set_update_fields(doc: Document) -> None:
    settings = doc.settings.element
    found = settings.find(qn("w:updateFields"))
    if found is None:
        update_fields = OxmlElement("w:updateFields")
        update_fields.set(qn("w:val"), "true")
        settings.append(update_fields)
    else:
        found.set(qn("w:val"), "true")


def normalize_toc_instruction(instr: str, max_level: int = 3) -> str:
    if "\\o" in instr:
        return re.sub(r'\\o\s+"1-\d+"', lambda _: f'\\o "1-{max_level}"', instr)
    if "TOC" in instr:
        return instr.replace("TOC", f'TOC \\o "1-{max_level}"', 1)
    return instr


def update_toc_fields(doc: Document, max_level: int = 3) -> None:
    for element in doc._element.iter():
        if element.tag.endswith("instrText") and element.text and "TOC" in element.text:
            element.text = normalize_toc_instruction(element.text, max_level=max_level)
        elif element.tag.endswith("fldSimple"):
            instr = element.get(qn("w:instr"))
            if instr and "TOC" in instr:
                element.set(qn("w:instr"), normalize_toc_instruction(instr, max_level=max_level))


def set_section_layout(doc: Document) -> None:
    for section in doc.sections:
        section.page_width = 7560310
        section.page_height = 10692130
        section.top_margin = 720090
        section.bottom_margin = 720090
        section.left_margin = 899795
        section.right_margin = 720090


def set_run_font(
    run,
    *,
    east_asia: str | None = "宋体",
    latin: str | None = "Times New Roman",
    size: float | None = None,
    bold: bool | None = None,
) -> None:
    if latin is not None:
        run.font.name = latin
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    if east_asia is not None:
        rfonts.set(qn("w:eastAsia"), east_asia)
    if latin is not None:
        rfonts.set(qn("w:ascii"), latin)
        rfonts.set(qn("w:hAnsi"), latin)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def set_paragraph_format(
    paragraph,
    *,
    alignment=None,
    first_line_indent_pt: float | None = None,
    line_spacing: float | None = None,
    line_spacing_pt: float | None = None,
    space_before_pt: float | None = None,
    space_after_pt: float | None = None,
) -> None:
    if alignment is not None:
        paragraph.alignment = alignment
    fmt = paragraph.paragraph_format
    fmt.first_line_indent = Pt(first_line_indent_pt) if first_line_indent_pt is not None else None
    if line_spacing_pt is not None:
        fmt.line_spacing = Pt(line_spacing_pt)
    elif line_spacing is not None:
        fmt.line_spacing = line_spacing
    else:
        fmt.line_spacing = None
    fmt.space_before = Pt(space_before_pt) if space_before_pt is not None else None
    fmt.space_after = Pt(space_after_pt) if space_after_pt is not None else None


def clear_runs(paragraph) -> None:
    for run in list(paragraph.runs):
        paragraph._p.remove(run._element)


def rewrite_paragraph(paragraph, segments: list[tuple[str, dict]]) -> None:
    clear_runs(paragraph)
    for text, style in segments:
        run = paragraph.add_run(text)
        set_run_font(run, **style)


def apply_font_to_paragraph(
    paragraph,
    *,
    east_asia: str = "宋体",
    latin: str = "Times New Roman",
    size: float | None = None,
    bold: bool | None = None,
) -> None:
    for run in paragraph.runs:
        set_run_font(run, east_asia=east_asia, latin=latin, size=size, bold=bold)


def first_nonempty_index(doc: Document, start: int = 0) -> int | None:
    for idx in range(start, len(doc.paragraphs)):
        if doc.paragraphs[idx].text.strip():
            return idx
    return None


def next_nonempty_index(doc: Document, start: int) -> int | None:
    return first_nonempty_index(doc, start)


def normalize_text(text: str) -> str:
    return text.replace(" ", "").replace("\t", "")


def format_cover(doc: Document) -> None:
    for idx, paragraph in enumerate(doc.paragraphs[:20]):
        text = paragraph.text.strip()
        compact = normalize_text(text)
        if not text:
            continue

        if text == "本 科 生 毕 业 设 计 说 明 书":
            set_paragraph_format(paragraph, alignment=WD_ALIGN_PARAGRAPH.CENTER)
            apply_font_to_paragraph(paragraph, east_asia="宋体", size=26, bold=True)
            continue

        if idx < 20 and text.startswith("基于深度学习"):
            set_paragraph_format(paragraph, alignment=WD_ALIGN_PARAGRAPH.CENTER)
            apply_font_to_paragraph(paragraph, east_asia="楷体", size=18, bold=True)
            continue

        if idx < 20 and text == "陈刚":
            set_paragraph_format(
                paragraph,
                alignment=WD_ALIGN_PARAGRAPH.CENTER,
                first_line_indent_pt=36.15,
            )
            apply_font_to_paragraph(paragraph, east_asia="楷体", size=18, bold=True)
            continue

        if compact.startswith(("学院：", "专业：", "班级：", "学号：", "指导教师：", "职称：")):
            set_paragraph_format(paragraph, first_line_indent_pt=71.5, line_spacing=1.5)
            for run in paragraph.runs:
                is_label = "：" in run.text or ":" in run.text
                set_run_font(
                    run,
                    east_asia="楷体_GB2312",
                    size=16,
                    bold=True if is_label else True,
                )
            continue

        if idx < 20 and ("年" in text and "月" in text):
            set_paragraph_format(paragraph, alignment=WD_ALIGN_PARAGRAPH.CENTER)
            apply_font_to_paragraph(paragraph, east_asia="楷体_GB2312", size=16, bold=True)


def format_statements(doc: Document) -> None:
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue

        if text in {"原创性声明", "指导声明"}:
            set_paragraph_format(paragraph, alignment=WD_ALIGN_PARAGRAPH.CENTER)
            apply_font_to_paragraph(paragraph, east_asia="黑体", size=22, bold=None)
            continue

        if text.startswith("本人郑重声明"):
            set_paragraph_format(paragraph, first_line_indent_pt=32)
            apply_font_to_paragraph(paragraph, east_asia="仿宋_GB2312", size=16, bold=None)
            continue

        if text.startswith("本人指导的"):
            set_paragraph_format(paragraph, first_line_indent_pt=21.85, line_spacing=1.5)
            apply_font_to_paragraph(paragraph, east_asia="仿宋_GB2312", size=16, bold=None)
            continue

        if text.startswith(("学生签名", "指导教师签名")):
            set_paragraph_format(paragraph, alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.5)
            apply_font_to_paragraph(paragraph, east_asia="仿宋_GB2312", size=16, bold=None)


def rewrite_label_body_paragraph(
    paragraph,
    *,
    label: str,
    body: str,
    label_face: str,
    body_face: str,
    size: float = 12,
    body_latin: str = "Times New Roman",
) -> None:
    rewrite_paragraph(
        paragraph,
        [
            (label, {"east_asia": label_face, "latin": "Times New Roman", "size": size, "bold": True}),
            (body, {"east_asia": body_face, "latin": body_latin, "size": size, "bold": False}),
        ],
    )


def format_title_and_abstract(doc: Document) -> tuple[int, int]:
    en_title_idx = next(
        (idx for idx, p in enumerate(doc.paragraphs) if p.text.strip().startswith(EN_TITLE_PREFIX)),
        None,
    )
    if en_title_idx is None:
        raise ValueError("未定位到英文题名页。")

    cn_title_idx = None
    for idx in range(en_title_idx - 1, -1, -1):
        text = doc.paragraphs[idx].text.strip()
        if text.startswith("基于深度学习"):
            cn_title_idx = idx
            break
    if cn_title_idx is None:
        raise ValueError("未定位到中文题名页。")

    # 中文题名页
    cn_name_idx = next_nonempty_index(doc, cn_title_idx + 1)
    cn_info_idx = next_nonempty_index(doc, (cn_name_idx or cn_title_idx) + 1)

    cn_title = doc.paragraphs[cn_title_idx]
    cn_title.style = doc.styles["Normal"]
    set_paragraph_format(cn_title, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    apply_font_to_paragraph(cn_title, east_asia="黑体", size=22, bold=True)

    if cn_name_idx is not None:
        paragraph = doc.paragraphs[cn_name_idx]
        paragraph.style = doc.styles["Normal"]
        set_paragraph_format(paragraph, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        apply_font_to_paragraph(paragraph, east_asia="宋体", size=14, bold=None)

    if cn_info_idx is not None:
        paragraph = doc.paragraphs[cn_info_idx]
        paragraph.style = doc.styles["Normal"]
        set_paragraph_format(paragraph, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        apply_font_to_paragraph(paragraph, east_asia="楷体_GB2312", size=14, bold=None)

    cn_kw_idx = next(
        (
            idx
            for idx in range(cn_title_idx + 1, en_title_idx)
            if doc.paragraphs[idx].text.strip().startswith("关键词")
        ),
        None,
    )
    if cn_kw_idx is None:
        raise ValueError("未定位到中文关键词段。")

    cn_abstract_indices: list[int] = []
    for idx in range((cn_info_idx or cn_title_idx) + 1, cn_kw_idx):
        text = doc.paragraphs[idx].text.strip()
        if text:
            cn_abstract_indices.append(idx)

    if cn_abstract_indices:
        first_idx = cn_abstract_indices[0]
        first_text = doc.paragraphs[first_idx].text.strip()
        if first_text.startswith("摘要："):
            first_body = first_text[len("摘要：") :]
        else:
            first_body = first_text
        rewrite_label_body_paragraph(
            doc.paragraphs[first_idx],
            label="摘要：",
            body=first_body,
            label_face="黑体",
            body_face="楷体_GB2312",
            size=12,
        )
        doc.paragraphs[first_idx].style = doc.styles["Normal"]
        set_paragraph_format(doc.paragraphs[first_idx])
        doc.paragraphs[first_idx].alignment = None

        for idx in cn_abstract_indices[1:]:
            paragraph = doc.paragraphs[idx]
            paragraph.style = doc.styles["Normal"]
            set_paragraph_format(paragraph)
            paragraph.alignment = None
            apply_font_to_paragraph(paragraph, east_asia="楷体_GB2312", size=12, bold=None)

    cn_kw_text = doc.paragraphs[cn_kw_idx].text.strip()
    cn_kw_value = cn_kw_text.split("：", 1)[1] if "：" in cn_kw_text else cn_kw_text.replace("关键词", "", 1)
    doc.paragraphs[cn_kw_idx].style = doc.styles["Normal"]
    rewrite_paragraph(
        doc.paragraphs[cn_kw_idx],
        [
            ("关键词", {"east_asia": "黑体", "latin": "Times New Roman", "size": 12, "bold": True}),
            ("：" + cn_kw_value, {"east_asia": "楷体_GB2312", "latin": "Times New Roman", "size": 12, "bold": False}),
        ],
    )
    doc.paragraphs[cn_kw_idx].alignment = None

    # 英文题名页
    en_name_idx = next_nonempty_index(doc, en_title_idx + 1)
    en_info_idx = next_nonempty_index(doc, (en_name_idx or en_title_idx) + 1)
    en_kw_idx = next(
        (
            idx
            for idx in range((en_info_idx or en_title_idx) + 1, len(doc.paragraphs))
            if doc.paragraphs[idx].text.strip().startswith("Keywords:")
        ),
        None,
    )
    if en_kw_idx is None:
        raise ValueError("未定位到英文关键词段。")

    doc.paragraphs[en_title_idx].style = doc.styles["Normal"]
    set_paragraph_format(doc.paragraphs[en_title_idx], alignment=WD_ALIGN_PARAGRAPH.CENTER)
    apply_font_to_paragraph(doc.paragraphs[en_title_idx], east_asia="黑体", size=14, bold=None)

    if en_name_idx is not None:
        doc.paragraphs[en_name_idx].style = doc.styles["Normal"]
        set_paragraph_format(doc.paragraphs[en_name_idx], alignment=WD_ALIGN_PARAGRAPH.CENTER)
        apply_font_to_paragraph(doc.paragraphs[en_name_idx], east_asia="楷体_GB2312", size=14, bold=None)

    if en_info_idx is not None:
        doc.paragraphs[en_info_idx].style = doc.styles["Normal"]
        set_paragraph_format(doc.paragraphs[en_info_idx], alignment=WD_ALIGN_PARAGRAPH.CENTER)
        apply_font_to_paragraph(doc.paragraphs[en_info_idx], east_asia="黑体", size=14, bold=None)

    en_abstract_indices: list[int] = []
    for idx in range((en_info_idx or en_title_idx) + 1, en_kw_idx):
        text = doc.paragraphs[idx].text.strip()
        if text:
            en_abstract_indices.append(idx)

    if en_abstract_indices:
        first_idx = en_abstract_indices[0]
        first_text = doc.paragraphs[first_idx].text.strip()
        if first_text.startswith("Abstract:"):
            first_body = first_text[len("Abstract:") :]
        else:
            first_body = first_text
        rewrite_paragraph(
            doc.paragraphs[first_idx],
            [
                ("Abstract:", {"east_asia": "黑体", "latin": "Times New Roman", "size": 12, "bold": True}),
                (first_body, {"east_asia": "黑体", "latin": "Times New Roman", "size": 12, "bold": False}),
            ],
        )
        doc.paragraphs[first_idx].style = doc.styles["Normal"]
        set_paragraph_format(doc.paragraphs[first_idx])
        doc.paragraphs[first_idx].alignment = None
        for idx in en_abstract_indices[1:]:
            doc.paragraphs[idx].style = doc.styles["Normal"]
            set_paragraph_format(doc.paragraphs[idx])
            doc.paragraphs[idx].alignment = None
            apply_font_to_paragraph(doc.paragraphs[idx], east_asia="黑体", size=12, bold=None)

    en_kw_text = doc.paragraphs[en_kw_idx].text.strip()
    en_kw_value = en_kw_text.split(":", 1)[1] if ":" in en_kw_text else en_kw_text.replace("Keywords", "", 1)
    doc.paragraphs[en_kw_idx].style = doc.styles["Normal"]
    rewrite_paragraph(
        doc.paragraphs[en_kw_idx],
        [
            ("Keywords:", {"east_asia": "黑体", "latin": "Times New Roman", "size": 12, "bold": True}),
            (en_kw_value, {"east_asia": "黑体", "latin": "Times New Roman", "size": 12, "bold": False}),
        ],
    )
    set_paragraph_format(doc.paragraphs[en_kw_idx])
    doc.paragraphs[en_kw_idx].alignment = None

    body_start_idx = next_nonempty_index(doc, en_kw_idx + 1)
    if body_start_idx is None:
        raise ValueError("未定位到正文起始段。")
    return body_start_idx, en_kw_idx


def normalize_chapter_heading(paragraph) -> None:
    text = paragraph.text.strip()
    match = CHAPTER_RE.match(text)
    if match:
        chapter_no, title = match.groups()
        clear_runs(paragraph)
        paragraph.add_run(f"{chapter_no} {title.strip()}")


def is_chapter_heading(text: str) -> bool:
    return bool(CHAPTER_RE.match(text) or (NORMALIZED_CHAPTER_RE.match(text) and "." not in text.split()[0]))


def is_section_heading(text: str) -> bool:
    return bool(SECTION_RE.match(text)) and not SUBSECTION_RE.match(text)


def is_subsection_heading(text: str) -> bool:
    return bool(SUBSECTION_RE.match(text))


def is_caption(text: str) -> bool:
    return bool(CAPTION_RE.match(text))


def strip_reference_number(text: str) -> str:
    return re.sub(r"^\[\d+\]\s*", "", text).strip()


def format_body(doc: Document, body_start_idx: int) -> None:
    ref_idx = next(
        (idx for idx, paragraph in enumerate(doc.paragraphs) if paragraph.text.strip() == REFERENCES_TITLE),
        len(doc.paragraphs),
    )
    appendix_idx = next(
        (idx for idx, paragraph in enumerate(doc.paragraphs) if APPENDIX_RE.match(paragraph.text.strip())),
        len(doc.paragraphs),
    )
    body_end_idx = min(ref_idx, appendix_idx)

    for idx in range(body_start_idx, body_end_idx):
        paragraph = doc.paragraphs[idx]
        text = paragraph.text.strip()
        if not text:
            continue

        normalize_chapter_heading(paragraph)
        text = paragraph.text.strip()

        if is_chapter_heading(text):
            try:
                paragraph.style = doc.styles["Heading 1"]
            except KeyError:
                paragraph.style = doc.styles["Normal"]
            set_paragraph_format(
                paragraph,
                first_line_indent_pt=6,
                line_spacing=1.0,
                space_before_pt=6,
                space_after_pt=6,
            )
            apply_font_to_paragraph(paragraph, east_asia="黑体", size=14, bold=False)
            continue

        if is_subsection_heading(text):
            try:
                paragraph.style = doc.styles["Heading 3"]
            except KeyError:
                paragraph.style = doc.styles["Normal"]
            set_paragraph_format(
                paragraph,
                first_line_indent_pt=5.95,
                line_spacing=1.0,
                space_before_pt=6,
                space_after_pt=6,
            )
            apply_font_to_paragraph(paragraph, east_asia="黑体", size=10.5, bold=None)
            continue

        if is_section_heading(text):
            try:
                paragraph.style = doc.styles["Heading 2"]
            except KeyError:
                paragraph.style = doc.styles["Normal"]
            set_paragraph_format(
                paragraph,
                first_line_indent_pt=5.95,
                line_spacing=1.0,
                space_before_pt=6,
                space_after_pt=6,
            )
            apply_font_to_paragraph(paragraph, east_asia="黑体", size=12, bold=None)
            continue

        if is_caption(text):
            paragraph.style = doc.styles["Normal"]
            set_paragraph_format(
                paragraph,
                alignment=WD_ALIGN_PARAGRAPH.CENTER,
                first_line_indent_pt=0,
                line_spacing=1.0,
            )
            apply_font_to_paragraph(paragraph, east_asia="宋体", size=9, bold=None)
            continue

        paragraph.style = doc.styles["Normal"]
        set_paragraph_format(
            paragraph,
            alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
            first_line_indent_pt=21,
            line_spacing_pt=20,
        )
        apply_font_to_paragraph(paragraph, east_asia="宋体", size=10.5, bold=False)

    if appendix_idx < len(doc.paragraphs):
        paragraph = doc.paragraphs[appendix_idx]
        paragraph.style = doc.styles["Heading 1"] if "Heading 1" in [s.name for s in doc.styles] else doc.styles["Normal"]
        set_paragraph_format(
            paragraph,
            first_line_indent_pt=6,
            line_spacing=1.0,
            space_before_pt=6,
            space_after_pt=6,
        )
        apply_font_to_paragraph(paragraph, east_asia="黑体", size=14, bold=False)


def format_references(doc: Document) -> None:
    ref_idx = next(
        (idx for idx, paragraph in enumerate(doc.paragraphs) if paragraph.text.strip() == REFERENCES_TITLE),
        None,
    )
    if ref_idx is None:
        return

    appendix_idx = next(
        (idx for idx, paragraph in enumerate(doc.paragraphs) if APPENDIX_RE.match(paragraph.text.strip())),
        len(doc.paragraphs),
    )

    title_paragraph = doc.paragraphs[ref_idx]
    title_paragraph.style = doc.styles["Normal"]
    set_paragraph_format(title_paragraph)
    apply_font_to_paragraph(title_paragraph, east_asia="宋体", size=12, bold=False)

    ref_no = 1
    for idx in range(ref_idx + 1, appendix_idx):
        paragraph = doc.paragraphs[idx]
        text = paragraph.text.strip()
        if not text:
            continue
        normalized_text = strip_reference_number(text)
        paragraph.style = doc.styles["Normal"]
        set_paragraph_format(paragraph)
        rewrite_paragraph(
            paragraph,
            [
                (
                    f"[{ref_no}] {normalized_text}",
                    {
                        "east_asia": "方正楷体_GB2312",
                        "latin": "Times New Roman",
                        "size": 9,
                        "bold": False,
                    },
                )
            ],
        )
        apply_font_to_paragraph(paragraph, east_asia="方正楷体_GB2312", size=9, bold=False)
        ref_no += 1


def format_tables(doc: Document) -> None:
    for table in doc.tables:
        for row_idx, row in enumerate(table.rows):
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if not paragraph.text.strip():
                        continue
                    set_paragraph_format(
                        paragraph,
                        alignment=WD_ALIGN_PARAGRAPH.CENTER,
                        line_spacing=1.0,
                    )
                    apply_font_to_paragraph(
                        paragraph,
                        east_asia="宋体",
                        size=10.5,
                        bold=True if row_idx == 0 else False,
                    )


def align_docx(target: Path, output: Path | None, in_place: bool) -> tuple[Path | None, Path, Path]:
    backup = backup_docx(target)
    doc = Document(target)
    set_update_fields(doc)
    update_toc_fields(doc, max_level=3)
    set_section_layout(doc)
    format_cover(doc)
    format_statements(doc)
    body_start_idx, _ = format_title_and_abstract(doc)
    format_body(doc, body_start_idx)
    format_references(doc)
    format_tables(doc)

    saved_target: Path | None = None
    if output is not None:
        ensure_output_parent(output)
        doc.save(output)

    if in_place:
        try:
            doc.save(target)
            saved_target = target
        except PermissionError:
            saved_target = None

    return saved_target, output or target, backup


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Align thesis DOCX format to the local sample style.")
    parser.add_argument("--target", type=Path, default=None, help="Target thesis DOCX path.")
    parser.add_argument("--output", type=Path, default=DEFAULT_OUTPUT, help="Write an aligned copy to this path.")
    parser.add_argument("--in-place", action="store_true", help="Also try to overwrite the target file.")
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    target = args.target if args.target is not None else find_target_docx()
    saved_target, output, backup = align_docx(target=target, output=args.output, in_place=args.in_place)
    print(f"Backup saved to: {backup}")
    print(f"Aligned copy saved to: {output}")
    if saved_target is None and args.in_place:
        print("Target file is likely open in Word, so the script kept the aligned copy only.")
    elif saved_target is not None:
        print(f"Target file updated in place: {saved_target}")


if __name__ == "__main__":
    main()
