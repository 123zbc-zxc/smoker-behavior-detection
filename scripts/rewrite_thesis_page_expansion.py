from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[1]
BACKUP_PATH = ROOT / "output" / "doc" / "thesis_backup_before_page_expansion.docx"
ALIGNED_COPY_PATH = ROOT / "output" / "doc" / "thesis_sample_aligned.docx"
EXPANDED_COPY_PATH = ROOT / "output" / "doc" / "thesis_page_expanded_working.docx"


INSERTIONS: dict[int, list[str]] = {
    170: [
        "如果进一步从目标形态上看，cigarette 类之所以困难，并不只是因为目标小。它往往还带有细长、低对比度和局部遮挡三个特征。对模型来说，这类目标在下采样之后剩下的有效信息本来就不多，一旦再叠加手部遮挡或压缩噪声，检测头能够利用的判别线索就会明显减少。",
        "另一个容易被忽略的问题是标注难度。香烟本体长度有限，姿态变化又大，不同样本中框选尺度很难完全一致。对于大目标来说，这类标注波动的影响通常不算明显；但对 cigarette 这类小目标来说，少量框偏差就可能直接影响 IoU 计算和训练稳定性，因此它比 smoke 和 smoking_person 更容易在指标上表现出波动。",
    ],
    178: [
        "还需要看到，smoke_bal 的主要作用是把实验成本压到当前算力能够承受的范围内，而不是专门为 ECA 这类结构改进量身定制的数据集。换句话说，当前数据组织方式首先服务于“把链路跑通并完成对比”，并没有对注意力模块做更细的样本强化，因此结构收益没有被放大也是可以理解的。",
    ],
    185: [
        "从案例复核的角度说，当前系统面对困难样本时呈现出的并不是随机失效，而是具有较稳定的失误模式：远距离场景更容易漏掉 cigarette，弱光和逆光条件下 smoke 边界更难判断，多人交叉时 smoking_person 的框位置更容易抖动。把这些问题拆开看，有助于后续有针对性地补数据，而不是笼统地归结为“模型效果不够好”。",
    ],
    189: [
        "也正因为有了这一轮完整对比，本文在后续系统实现中才优先选择 baseline 作为默认权重，而不是继续把 ECA 版本放在演示主链路上。这样做并不是否定改进实验本身，而是让系统展示优先建立在当前更稳定的模型结果之上，使论文中的实验结论和答辩演示保持一致。",
    ],
    193: [
        "因此，本章的分析更适合被理解为一组阶段性结论：当前方案已经可以覆盖典型场景，并具备继续优化的基础，但还没有达到可以脱离场景条件讨论泛化能力的程度。这样的表述既保留了研究结果，也避免把模型能力说得过满。",
    ],
    200: [
        "如果按功能分层来看，当前系统大致可以分成四层：第一层是数据与训练层，负责生成数据集、训练权重和实验摘要；第二层是推理服务层，负责读取模型、接收输入并产出检测结果；第三层是网页展示层，负责把模型状态、指标摘要和案例结果组织到浏览器界面；第四层是记录存储层，负责把模型配置、检测记录和任务状态长期保存下来。这样分层后，训练代码和展示代码虽然相互配合，但职责并没有混在一起。",
    ],
    215: [
        "从数据流角度看，网页端的价值不仅在于“能看到结果”，还在于它把输入、推理、落盘和记录查询串成了一条完整路径。老师在答辩现场看到的页面结果，并不是临时拼出来的截图，而是系统真实运行后生成的检测记录、标注图和数据库信息。这一点能够明显增强系统部分的可信度。",
    ],
    221: [
        "这种模块划分还有一个实际好处，就是修改范围更容易控制。比如更换默认模型时，通常只需要调整配置和模型注册信息；如果要补充新的网页接口，主要影响的是路由层和服务层；若要改进训练流程，则更多集中在 scripts 和 models 目录。对于毕业设计阶段的项目来说，这种结构已经足以支撑后续继续完善。",
    ],
    225: [
        "数据库之所以有必要保留下来，核心原因在于网页系统并不是只做一次性的前端展示。没有记录层，图片检测做完之后就只剩下一张标注图，后续很难追溯使用的是哪一个模型、对应哪些检测框、任务何时执行完成。把这些信息结构化存下来之后，论文截图整理、案例回查和答辩演示都会方便得多。",
    ],
    234: [
        "稳定性设计里还有一个比较实际的考虑，就是尽量减少答辩现场的不可控因素。当前系统把权重路径、结果目录、数据库状态和接口健康检查放在同一套启动流程里，本质上是在把常见故障尽量前置暴露。这样做虽然谈不上复杂，但比等到上传文件或切换页面时才发现问题要稳妥得多。",
    ],
    238: [
        "从测试组织方式看，人工核验和自动脚本各自承担的角色也比较明确。人工核验更适合检查界面展示是否顺畅、结果是否便于说明；自动冒烟测试则更适合反复检查接口是否还能正常返回、数据库是否还能写入、视频任务状态是否还能更新。两者结合之后，系统测试就不再只是展示性操作，而是具备了一定的重复验证能力。",
    ],
    241: [
        "需要说明的是，这里所说的“满足答辩所需稳定性”，仍然是面向毕业设计原型系统的判断标准。它强调的是链路清楚、结果可复现、页面可展示，而不是面向高并发、长周期运行或复杂权限体系的工程化要求。把这一点交代清楚，反而能让系统部分的结论更加可信。",
    ],
}


def find_target_docx() -> Path:
    matches = [
        p
        for p in ROOT.glob("*.docx")
        if "224-" in p.name and not p.name.startswith("~$") and p.stat().st_size > 1000
    ]
    if not matches:
        raise FileNotFoundError("未找到当前定稿论文 DOCX 文件。")
    matches.sort()
    return matches[0]


def backup_docx(target: Path) -> None:
    BACKUP_PATH.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(target, BACKUP_PATH)


def insert_paragraph_after(paragraph: Paragraph, text: str, style_name: str = "Normal") -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    new_para.style = style_name
    new_para.add_run(text)
    return new_para


def build_expanded_doc(source_path: Path) -> Document:
    doc = Document(source_path)
    anchors = {idx: doc.paragraphs[idx] for idx in INSERTIONS}
    for idx in sorted(INSERTIONS):
        anchor = anchors[idx]
        current = anchor
        for text in INSERTIONS[idx]:
            current = insert_paragraph_after(current, text)
    return doc


def main() -> None:
    target = find_target_docx()
    backup_docx(target)
    doc = build_expanded_doc(target)
    save_mode = "in_place"
    try:
        doc.save(target)
    except PermissionError:
        EXPANDED_COPY_PATH.parent.mkdir(parents=True, exist_ok=True)
        doc.save(EXPANDED_COPY_PATH)
        save_mode = "copy_only"
    if save_mode == "in_place" and ALIGNED_COPY_PATH.exists():
        shutil.copy2(target, ALIGNED_COPY_PATH)
    elif save_mode == "copy_only":
        shutil.copy2(EXPANDED_COPY_PATH, ALIGNED_COPY_PATH)
    print(f"Backup saved to: {BACKUP_PATH}")
    if save_mode == "in_place":
        print(f"Expanded thesis saved to: {target}")
    else:
        print(f"Target file is locked, expanded copy saved to: {EXPANDED_COPY_PATH}")
    if ALIGNED_COPY_PATH.exists():
        print(f"Aligned copy synced to: {ALIGNED_COPY_PATH}")


if __name__ == "__main__":
    main()
