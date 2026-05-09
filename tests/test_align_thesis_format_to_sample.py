from __future__ import annotations

import importlib.util
from pathlib import Path
import unittest

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH


ROOT = Path(__file__).resolve().parents[1]
MODULE_PATH = ROOT / "scripts" / "align_thesis_format_to_sample.py"
SPEC = importlib.util.spec_from_file_location("align_thesis_format_to_sample", MODULE_PATH)
MODULE = importlib.util.module_from_spec(SPEC)
assert SPEC.loader is not None
SPEC.loader.exec_module(MODULE)


def east_asia_font(run) -> str | None:
    rpr = run._element.rPr
    if rpr is None or rpr.rFonts is None:
        return None
    return rpr.rFonts.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia")


class AlignThesisFormatTests(unittest.TestCase):
    def test_normalize_toc_instruction_expands_to_three_levels(self) -> None:
        self.assertEqual(
            MODULE.normalize_toc_instruction('TOC \\o "1-2" \\h \\u'),
            'TOC \\o "1-3" \\h \\u',
        )

    def test_format_body_promotes_section_levels_into_toc_heading_styles(self) -> None:
        doc = Document()
        doc.add_paragraph("1 绪论")
        doc.add_paragraph("1.1 研究背景")
        doc.add_paragraph("1.1.1 研究意义")
        doc.add_paragraph("这里是正文。")
        doc.add_paragraph("参考文献")

        MODULE.format_body(doc, 0)

        self.assertEqual(doc.paragraphs[0].style.name, "Heading 1")
        self.assertEqual(doc.paragraphs[1].style.name, "Heading 2")
        self.assertEqual(doc.paragraphs[2].style.name, "Heading 3")

    def test_is_caption_only_matches_real_caption_lines(self) -> None:
        self.assertTrue(MODULE.is_caption("图4-1 系统总体架构图"))
        self.assertTrue(MODULE.is_caption("表3-6 Baseline训练参数设置"))
        self.assertFalse(MODULE.is_caption("表3-6给出了 baseline 训练方案的主要参数设置。"))

    def test_format_references_adds_square_bracket_numbers(self) -> None:
        doc = Document()
        doc.add_paragraph("参考文献")
        doc.add_paragraph("Redmon J, Divvala S, Girshick R, et al. You Only Look Once.")
        doc.add_paragraph("[2] Wang Q, Wu B, Zhu P, et al. ECA-Net.")
        doc.add_paragraph("附录：项目对应的关键命令")

        MODULE.format_references(doc)

        self.assertEqual(doc.paragraphs[1].text, "[1] Redmon J, Divvala S, Girshick R, et al. You Only Look Once.")
        self.assertEqual(doc.paragraphs[2].text, "[2] Wang Q, Wu B, Zhu P, et al. ECA-Net.")
        self.assertEqual(east_asia_font(doc.paragraphs[1].runs[0]), "方正楷体_GB2312")

    def test_format_tables_sets_header_bold_and_centers_cells(self) -> None:
        doc = Document()
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "数据来源"
        table.cell(0, 1).text = "图像数量"
        table.cell(1, 0).text = "AI Studio VOC 数据"
        table.cell(1, 1).text = "783"

        MODULE.format_tables(doc)

        header_paragraph = table.cell(0, 0).paragraphs[0]
        body_paragraph = table.cell(1, 0).paragraphs[0]

        self.assertEqual(header_paragraph.alignment, WD_ALIGN_PARAGRAPH.CENTER)
        self.assertTrue(all(run.bold for run in header_paragraph.runs if run.text))
        self.assertEqual(header_paragraph.runs[0].font.size.pt, 10.5)
        self.assertEqual(body_paragraph.alignment, WD_ALIGN_PARAGRAPH.CENTER)
        self.assertTrue(all(run.bold is False for run in body_paragraph.runs if run.text))


if __name__ == "__main__":
    unittest.main()
